from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path

from openpyxl import load_workbook
from pypdf import PdfReader

from rag.knowledge_base import _candidate_files, _tokens
from rag.ocr import is_image_file, extract_image_text
from utils.config import load_settings, safe_workspace_id
from llm.ollama_client import generate_response


COUNT_RE = re.compile(r"\b(how\s+many|count|number\s+of|total)\b", re.I)
PEOPLE_RE = re.compile(r"\b(people|persons?|students?|candidates?|eligible|names?)\b", re.I)
QUESTION_RE = re.compile(r"\b(questions?|problems?|coding\s+questions?)\b", re.I)
PDF_RE = re.compile(r"\b(pdf|document|file|list)\b", re.I)
REFERENCE_ID_RE = re.compile(r"\b[A-Z]{1,4}\d{6,}\b")
DOCUMENT_RE = re.compile(r"\b(pdf|document|file|list|sheet|spreadsheet|docx|word|excel|xlsx|uploaded|last|this)\b", re.I)
MAX_FULL_CONTEXT_CHARS = 90000


@dataclass(frozen=True)
class DocumentCount:
    file: Path
    label: str
    count: int
    unique_count: int | None = None
    method: str = ""


@dataclass(frozen=True)
class DocumentRecord:
    reference_id: str
    name: str


def _read_pdf(path: Path) -> str:
    reader = PdfReader(path)
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _text_for_file(path: Path) -> str:
    suffix = path.suffix.lower()
    if is_image_file(path):
        return extract_image_text(path)
    if suffix == ".pdf":
        return _read_pdf(path)
    if suffix in {".txt", ".md"}:
        return path.read_text(encoding="utf-8", errors="ignore")
    if suffix == ".docx":
        from docx import Document

        document = Document(path)
        parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if values:
                    parts.append(" | ".join(values))
        return "\n".join(parts)
    if suffix == ".xlsx":
        workbook = load_workbook(path, data_only=True)
        parts = []
        for sheet in workbook.worksheets:
            parts.append(f"Sheet: {sheet.title}")
            for row in sheet.iter_rows(values_only=True):
                values = [str(value).strip() for value in row if value not in (None, "")]
                if values:
                    parts.append(" | ".join(values))
        return "\n".join(parts)
    return ""


def _count_xlsx(path: Path) -> DocumentCount | None:
    workbook = load_workbook(path, data_only=True)
    total = 0
    for sheet in workbook.worksheets:
        rows = [
            row
            for row in sheet.iter_rows(values_only=True)
            if any(value not in (None, "") for value in row)
        ]
        if not rows:
            continue
        header_offset = 1 if any(isinstance(value, str) for value in rows[0]) else 0
        total += max(len(rows) - header_offset, 0)
    if total:
        return DocumentCount(path, "rows/people", total, method="non-empty spreadsheet rows excluding header rows")
    return None


def _count_text_records(path: Path) -> DocumentCount | None:
    text = _text_for_file(path)
    if not text.strip():
        return None

    ids = REFERENCE_ID_RE.findall(text)
    if ids:
        return DocumentCount(
            path,
            "candidate/person records",
            len(ids),
            unique_count=len(set(ids)),
            method="reference IDs in the full document",
        )

    lines = [
        line.strip()
        for line in text.splitlines()
        if line.strip() and not re.fullmatch(r"[\W_]+", line.strip())
    ]
    data_lines = [
        line
        for line in lines
        if not re.search(r"\b(reference|candidate|name|sr\.?\s*no|serial|page)\b", line, flags=re.I)
    ]
    if data_lines:
        return DocumentCount(path, "non-empty text lines", len(data_lines), method="non-empty data-like lines")
    return None


def _question_titles(text: str) -> list[str]:
    titles: list[str] = []
    seen: set[str] = set()
    patterns = [
        r"(?:^|\n)\s*(?:question|problem|q)\s*(?:number|no\.?|#)?\s*(\d{1,3})\s*[:.)-]\s*([^\n]{1,160})",
        r"(?:^|\n)\s*(\d{1,3})\s*[.)-]\s+([A-Za-z][^\n]{3,160})",
    ]
    for pattern in patterns:
        for match in re.finditer(pattern, text, flags=re.I):
            number = int(match.group(1))
            title = re.sub(r"\s+", " ", match.group(2)).strip(" :-")
            if not title or number < 1:
                continue
            key = f"{number}:{title.lower()}"
            if key not in seen:
                seen.add(key)
                titles.append(f"{number}. {title}")
    return titles


def _count_questions(path: Path) -> DocumentCount | None:
    text = _text_for_file(path)
    titles = _question_titles(text)
    if titles:
        return DocumentCount(path, "questions", len(titles), method="question headings")

    numbered = re.findall(r"(?:^|\n)\s*(\d{1,3})\s*[.)-]\s+", text)
    if numbered:
        return DocumentCount(path, "questions", len(set(numbered)), method="numbered question entries")
    return None


def _count_file(path: Path) -> DocumentCount | None:
    if path.suffix.lower() == ".xlsx":
        return _count_xlsx(path)
    if path.suffix.lower() in {".pdf", ".docx", ".txt", ".md"}:
        return _count_text_records(path)
    return None


def _records_from_text(text: str) -> list[DocumentRecord]:
    matches = list(REFERENCE_ID_RE.finditer(text))
    records: list[DocumentRecord] = []
    for index, match in enumerate(matches):
        start = match.end()
        end = matches[index + 1].start() if index + 1 < len(matches) else len(text)
        name = re.sub(r"\s+", " ", text[start:end]).strip(" |,-:")
        name = re.sub(r"\b(reference|candidate|name)\b", "", name, flags=re.I).strip(" |,-:")
        if name:
            records.append(DocumentRecord(match.group(0), name))
    return records


def _records_for_file(path: Path) -> list[DocumentRecord]:
    if path.suffix.lower() not in {".pdf", ".docx", ".txt", ".md", ".xlsx"}:
        return []
    return _records_from_text(_text_for_file(path))


def _requested_limit(query: str, default: int = 20) -> int:
    word_numbers = {
        "one": 1,
        "two": 2,
        "three": 3,
        "four": 4,
        "five": 5,
        "six": 6,
        "seven": 7,
        "eight": 8,
        "nine": 9,
        "ten": 10,
    }
    digit = re.search(r"\b(?:first|top|last)?\s*(\d{1,3})\b", query, flags=re.I)
    if digit:
        return max(1, min(int(digit.group(1)), 100))
    lower = query.lower()
    for word, value in word_numbers.items():
        if re.search(rf"\b(?:first|top|last)?\s*{word}\b", lower):
            return value
    if re.search(r"\b(all|complete|full)\b", lower):
        return 100
    return default


def answer_document_records(query: str, tenant_id: str | None = None, preferred_file: str | None = None) -> str | None:
    if not re.search(r"\b(list|show|give|first|last|top|details?|find|search|who|which)\b", query, flags=re.I):
        return None
    if not re.search(r"\b(candidate|student|person|people|name|eligible|reference|id|pdf|document|file|list)\b", query, flags=re.I):
        return None

    target = _target_file(query, tenant_id, preferred_file)
    if not target:
        return None

    try:
        records = _records_for_file(target)
    except Exception as exc:
        return f"I found {target.name}, but I could not read its text clearly. {exc}"
    if not records:
        return None

    lower = query.lower()
    selected = records
    id_match = REFERENCE_ID_RE.search(query)
    if id_match:
        selected = [record for record in records if record.reference_id.upper() == id_match.group(0).upper()]
    else:
        name_terms = [
            token
            for token in re.findall(r"[a-zA-Z]{3,}", query)
            if token.lower()
            not in {
                "list",
                "show",
                "give",
                "first",
                "last",
                "top",
                "details",
                "find",
                "search",
                "candidate",
                "candidates",
                "student",
                "students",
                "person",
                "people",
                "name",
                "names",
                "eligible",
                "reference",
                "from",
                "this",
                "that",
                "document",
                "file",
                "pdf",
                "tcs",
            }
        ]
        if name_terms:
            selected = [
                record
                for record in records
                if all(term.lower() in record.name.lower() for term in name_terms)
            ] or records

    if "last" in lower and not id_match:
        selected = list(reversed(selected))

    limit = _requested_limit(query)
    visible = selected[:limit]
    if not visible:
        return f"I found {len(records)} candidate/person records in {target.name}, but no matching record was found."

    lines = [
        f"I found {len(records)} candidate/person records in {target.name}.",
        f"Here are {len(visible)} record(s):",
    ]
    for index, record in enumerate(visible, start=1):
        lines.append(f"{index}. {record.reference_id} - {record.name}")
    if len(selected) > len(visible):
        lines.append(f"{len(selected) - len(visible)} more matching record(s) are available.")
    return "\n".join(lines)


def _target_file(query: str, tenant_id: str | None, preferred_file: str | None = None) -> Path | None:
    settings = load_settings()
    workspace = safe_workspace_id(tenant_id)
    allowed_roots = [settings.uploads_dir / workspace, settings.generated_dir / workspace]

    if preferred_file:
        path = Path(preferred_file)
        try:
            resolved = path.resolve()
        except OSError:
            resolved = path
        if path.exists() and any(root.resolve() in resolved.parents for root in allowed_roots if root.exists()):
            return path

    candidates = _candidate_files(settings, tenant_id)
    if not candidates:
        return None

    query_tokens = _tokens(query)
    scored: list[tuple[int, float, Path]] = []
    for path in candidates:
        name_tokens = _tokens(path.stem.replace("_", " "))
        score = len(query_tokens & name_tokens) * 4
        if path.suffix.lower() == ".pdf" and PDF_RE.search(query):
            score += 2
        if "eligible" in query_tokens and "eligible" in name_tokens:
            score += 8
        if score:
            scored.append((score, path.stat().st_mtime, path))

    if scored:
        scored.sort(reverse=True)
        return scored[0][2]
    return sorted(candidates, key=lambda item: item.stat().st_mtime, reverse=True)[0]


def target_document(query: str, tenant_id: str | None = None, preferred_file: str | None = None) -> Path | None:
    return _target_file(query, tenant_id, preferred_file)


def answer_document_count(query: str, tenant_id: str | None = None, preferred_file: str | None = None) -> str | None:
    if not (COUNT_RE.search(query) and (PEOPLE_RE.search(query) or QUESTION_RE.search(query) or PDF_RE.search(query))):
        return None

    target = _target_file(query, tenant_id, preferred_file)
    if not target:
        return None

    try:
        count = _count_questions(target) if QUESTION_RE.search(query) else _count_file(target)
    except Exception as exc:
        return f"I found {target.name}, but I could not read its text clearly. {exc}"
    if not count:
        return f"I found {target.name}, but I could not extract countable records from it."

    unique_note = ""
    if count.unique_count is not None and count.unique_count != count.count:
        unique_note = f" ({count.unique_count} unique reference IDs; duplicate IDs appear in the file)"

    label = "eligible people/students" if "eligible" in target.stem.lower() or re.search(r"\beligible\b", query, re.I) else count.label
    return f"{target.name} contains {count.count} {label}{unique_note}."


def answer_full_document_query(query: str, tenant_id: str | None = None, preferred_file: str | None = None) -> str | None:
    if wants_web_question(query):
        return None
    if not preferred_file and not DOCUMENT_RE.search(query):
        return None

    target = _target_file(query, tenant_id, preferred_file)
    if not target:
        return None

    try:
        text = _text_for_file(target)
    except Exception as exc:
        return f"I found {target.name}, but I could not read its text clearly. {exc}"
    if not text.strip():
        return f"I found {target.name}, but I could not extract readable text/table content from it."

    truncated = len(text) > MAX_FULL_CONTEXT_CHARS
    full_context = text[:MAX_FULL_CONTEXT_CHARS]
    instruction = (
        "Answer using the complete extracted document content supplied below. "
        "Do not answer from memory or partial excerpts. For lists, counts, names, IDs, statuses, dates, "
        "tables, eligibility, totals, clauses, and conditions, inspect all supplied content carefully. "
        "If the supplied content is truncated, say that the answer is based on the visible extracted content. "
        "Cite the file name and mention exact rows, IDs, names, clauses, or fields used when useful."
    )
    if truncated:
        instruction += " The extracted content was too large for one model request and has been truncated."

    return generate_response(
        f"{instruction}\n\nUser question:\n{query}",
        f"Full extracted content from {target.name}:\n{full_context}",
    )


def wants_web_question(query: str) -> bool:
    return bool(re.search(r"\b(web|internet|online|latest|current|today|news|search)\b", query, flags=re.I))
