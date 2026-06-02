from __future__ import annotations

import contextlib
import csv
import hashlib
import re
import sys
import types
from dataclasses import dataclass
from functools import lru_cache
from html.parser import HTMLParser
from pathlib import Path
from typing import Iterator

from pypdf import PdfReader

from rag.ocr import extract_image_text, is_image_file
from utils.config import BASE_DIR, load_settings, safe_workspace_id, tenant_generated_dir


CONVERTIBLE_EXTENSIONS = {
    ".pdf",
    ".doc",
    ".docx",
    ".ppt",
    ".pptx",
    ".xls",
    ".xlsx",
    ".csv",
    ".html",
    ".htm",
    ".txt",
    ".text",
    ".md",
    ".mhtml",
    ".xml",
    ".json",
    ".rtf",
    ".odt",
    ".epub",
    ".asciidoc",
    ".adoc",
} | {".jpg", ".jpeg", ".png", ".webp"}
MARKITDOWN_SRC = BASE_DIR / "markitdown-main" / "markitdown-main" / "packages" / "markitdown" / "src"
MARKDOWN_CACHE_DIRNAME = "_markdown_ingest"


@dataclass(frozen=True)
class MarkdownConversion:
    source_path: Path
    markdown_path: Path | None
    converter: str
    error: str = ""


@contextlib.contextmanager
def _temporary_sys_path(path: Path) -> Iterator[None]:
    value = str(path)
    added = path.exists() and value not in sys.path
    if added:
        sys.path.insert(0, value)
    try:
        yield
    finally:
        if added:
            with contextlib.suppress(ValueError):
                sys.path.remove(value)


def _markdown_dir(tenant_id: str | None) -> Path:
    settings = load_settings()
    root = tenant_generated_dir(settings, tenant_id)
    path = root / MARKDOWN_CACHE_DIRNAME
    path.mkdir(parents=True, exist_ok=True)
    return path


def is_markdown_sidecar(path: Path) -> bool:
    return path.suffix.lower() == ".md" and MARKDOWN_CACHE_DIRNAME in {part.lower() for part in path.parts}


def markdown_sidecar_path(path: Path, tenant_id: str | None) -> Path:
    workspace = safe_workspace_id(tenant_id)
    digest = hashlib.sha1(str(path.resolve()).lower().encode("utf-8", errors="ignore")).hexdigest()[:10]
    return _markdown_dir(workspace) / f"{path.stem}.{digest}.md"


def _write_markdown(path: Path, source: Path, markdown: str, converter: str) -> Path:
    source_label = source.name.replace("\n", " ").strip()
    body = markdown.strip()
    header = (
        f"<!-- AutoNova Markdown conversion\n"
        f"source: {source_label}\n"
        f"converter: {converter}\n"
        f"-->\n\n"
        f"# {source.stem}\n\n"
    )
    path.write_text(header + body + "\n", encoding="utf-8")
    return path


@lru_cache(maxsize=1)
def _docling_converter():
    from docling.datamodel.base_models import InputFormat
    from docling.datamodel.pipeline_options import PdfPipelineOptions
    from docling.document_converter import DocumentConverter, PdfFormatOption

    pdf_options = PdfPipelineOptions()
    pdf_options.do_ocr = False
    return DocumentConverter(
        format_options={
            InputFormat.PDF: PdfFormatOption(pipeline_options=pdf_options),
        }
    )


def _docling_convert(path: Path) -> str:
    result = _docling_converter().convert(str(path))
    document = getattr(result, "document", None)
    if document is None:
        return ""
    markdown = document.export_to_markdown()
    return markdown if isinstance(markdown, str) else str(markdown)


def _markitdown_convert(path: Path) -> str:
    with _temporary_sys_path(MARKITDOWN_SRC):
        if "magika" not in sys.modules:

            class _FallbackMagika:
                def identify_stream(self, _file_stream):
                    output = types.SimpleNamespace(label="unknown", is_text=False, extensions=[], mime_type=None)
                    prediction = types.SimpleNamespace(output=output)
                    return types.SimpleNamespace(status="unknown", prediction=prediction)

            sys.modules["magika"] = types.SimpleNamespace(Magika=_FallbackMagika)
        from markitdown import MarkItDown

        result = MarkItDown().convert(str(path))
        return getattr(result, "text_content", str(result))


def _docx_to_markdown(path: Path) -> str:
    from docx import Document

    document = Document(path)
    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table_index, table in enumerate(document.tables, start=1):
        rows = [[cell.text.strip().replace("\n", " ") for cell in row.cells] for row in table.rows]
        rows = [row for row in rows if any(row)]
        if not rows:
            continue
        width = max(len(row) for row in rows)
        normalized = [row + [""] * (width - len(row)) for row in rows]
        parts.append(f"\nTable {table_index}")
        parts.append("| " + " | ".join(normalized[0]) + " |")
        parts.append("| " + " | ".join("---" for _ in range(width)) + " |")
        for row in normalized[1:]:
            parts.append("| " + " | ".join(row) + " |")
    return "\n\n".join(parts)


def _pdf_to_markdown(path: Path) -> str:
    reader = PdfReader(path)
    parts: list[str] = []
    for index, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if text:
            parts.append(f"## Page {index}\n\n{text}")
    return "\n\n".join(parts)


def _xlsx_to_markdown(path: Path) -> str:
    from openpyxl import load_workbook

    workbook = load_workbook(path, data_only=True)
    parts: list[str] = []
    for sheet in workbook.worksheets:
        rows: list[list[str]] = []
        for row in sheet.iter_rows(values_only=True):
            values = ["" if value is None else str(value).strip().replace("\n", " ") for value in row]
            if any(values):
                rows.append(values)
        if not rows:
            continue
        width = max(len(row) for row in rows)
        normalized = [row + [""] * (width - len(row)) for row in rows]
        parts.append(f"## Sheet: {sheet.title}")
        parts.append("| " + " | ".join(normalized[0]) + " |")
        parts.append("| " + " | ".join("---" for _ in range(width)) + " |")
        for row in normalized[1:]:
            parts.append("| " + " | ".join(row) + " |")
    return "\n".join(parts)


def _csv_to_markdown(path: Path) -> str:
    with path.open("r", encoding="utf-8-sig", errors="ignore", newline="") as handle:
        rows = [[cell.strip() for cell in row] for row in csv.reader(handle)]
    rows = [row for row in rows if any(row)]
    if not rows:
        return ""
    width = max(len(row) for row in rows)
    normalized = [row + [""] * (width - len(row)) for row in rows]
    lines = ["| " + " | ".join(normalized[0]) + " |", "| " + " | ".join("---" for _ in range(width)) + " |"]
    lines.extend("| " + " | ".join(row) + " |" for row in normalized[1:])
    return "\n".join(lines)


class _TextHTMLParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.parts: list[str] = []

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        if tag in {"h1", "h2", "h3", "p", "li", "tr", "br"}:
            self.parts.append("\n")
        if tag == "li":
            self.parts.append("- ")

    def handle_data(self, data: str) -> None:
        text = re.sub(r"\s+", " ", data).strip()
        if text:
            self.parts.append(text + " ")

    def text(self) -> str:
        return re.sub(r"\n{3,}", "\n\n", "".join(self.parts)).strip()


def _html_to_markdown(path: Path) -> str:
    parser = _TextHTMLParser()
    parser.feed(path.read_text(encoding="utf-8", errors="ignore"))
    return parser.text()


def _fallback_convert(path: Path) -> str:
    suffix = path.suffix.lower()
    if is_image_file(path):
        return extract_image_text(path)
    if suffix == ".txt":
        return path.read_text(encoding="utf-8-sig", errors="ignore")
    if suffix == ".docx":
        return _docx_to_markdown(path)
    if suffix == ".pdf":
        return _pdf_to_markdown(path)
    if suffix == ".xlsx":
        return _xlsx_to_markdown(path)
    if suffix == ".csv":
        return _csv_to_markdown(path)
    if suffix in {".html", ".htm"}:
        return _html_to_markdown(path)
    return ""


def convert_to_markdown(path: Path, tenant_id: str | None) -> MarkdownConversion:
    suffix = path.suffix.lower()
    if suffix == ".md":
        return MarkdownConversion(path, path, "native")
    if suffix not in CONVERTIBLE_EXTENSIONS or not path.exists():
        return MarkdownConversion(path, None, "unsupported")

    output = markdown_sidecar_path(path, tenant_id)
    try:
        source_mtime = path.stat().st_mtime
        if output.exists() and output.stat().st_mtime >= source_mtime and output.stat().st_size > 0:
            return MarkdownConversion(path, output, "cached")
    except OSError:
        pass

    conversion_errors: list[str] = []
    try:
        markdown = _docling_convert(path)
        if markdown and markdown.strip():
            return MarkdownConversion(path, _write_markdown(output, path, markdown, "docling"), "docling")
    except Exception as exc:
        conversion_errors.append(f"docling: {exc}")

    try:
        markdown = _markitdown_convert(path)
        if markdown and markdown.strip():
            return MarkdownConversion(path, _write_markdown(output, path, markdown, "markitdown"), "markitdown")
    except Exception as exc:
        conversion_errors.append(f"markitdown: {exc}")

    try:
        markdown = _fallback_convert(path)
        if markdown and markdown.strip():
            return MarkdownConversion(
                path,
                _write_markdown(output, path, markdown, "autonova-fallback"),
                "autonova-fallback",
                "; ".join(conversion_errors),
            )
    except Exception as exc:
        conversion_errors.append(f"autonova-fallback: {exc}")
        return MarkdownConversion(path, None, "failed", "; ".join(conversion_errors))

    return MarkdownConversion(path, None, "empty", "; ".join(conversion_errors))


def readable_markdown_path(path: Path, tenant_id: str | None) -> Path:
    conversion = convert_to_markdown(path, tenant_id)
    if conversion.markdown_path and conversion.markdown_path.exists():
        return conversion.markdown_path
    return path


def readable_text(path: Path, tenant_id: str | None) -> str:
    source = readable_markdown_path(path, tenant_id)
    if source.suffix.lower() in {".txt", ".md"}:
        return source.read_text(encoding="utf-8-sig", errors="ignore")
    return ""
