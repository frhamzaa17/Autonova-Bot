from __future__ import annotations

import json
import re
from collections import Counter
from datetime import date, datetime, timedelta
from pathlib import Path
from typing import Any

from openpyxl import load_workbook
from pypdf import PdfReader

from rag.markdown_converter import CONVERTIBLE_EXTENSIONS, is_markdown_sidecar, readable_text
from rag.ocr import IMAGE_EXTENSIONS, extract_image_text, is_image_file
from utils.config import load_settings, safe_workspace_id


PROPERTY_ID_RE = re.compile(r"^(P[RC]-\d{3,})$", re.I)
AGREEMENT_ID_RE = re.compile(r"^([RC]A-\d{3,})$", re.I)
REFERENCE_ID_RE = re.compile(r"\b[A-Z]{1,4}\d{6,}\b")
QUESTION_TITLE_RE = re.compile(
    r"(?:^|\n)\s*(?:(?:question|problem|q)\s*(?:number|no\.?|#)?\s*)?(\d{1,3})\s*[:.)-]\s*([^\n]{3,180})",
    re.I,
)
MCQ_QUESTION_RE = re.compile(r"^\s*Q?(\d{1,3})\s*[\.)-]\s*(.+?)\s*$", re.I)
MCQ_OPTION_RE = re.compile(r"^\s*([A-D])\s*[\.)-]\s*(.+?)\s*$", re.I)
ANSWER_KEY_RE = re.compile(r"\b(\d{1,3})\s*[-:]\s*([A-D])\b", re.I)
SUPPORTED_TEXT_EXTENSIONS = {".md"} | CONVERTIBLE_EXTENSIONS | IMAGE_EXTENSIONS
STATUSES = ("Available", "Under Offer", "Sold", "Rented")
PHONE_RE = re.compile(r"\b(?:\d{2,5}[- ]?)?(?:\d|X){3,5}[- ]?(?:\d|X){3,5}(?:[- ]?(?:\d|X){2,5})?\b", re.I)
EMAIL_RE = re.compile(r"\b[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}\b", re.I)
DATE_RE = re.compile(r"\b\d{1,2}[-/][A-Za-z]{3}[-/]\d{2,4}\b|\b\d{1,2}[-/]\d{1,2}[-/]\d{2,4}\b", re.I)
AMOUNT_RE = re.compile(r"\b(?:rs\.?|inr)?\s*\d[\d,]*(?:\.\d+)?\s*(?:cr|crore|l|lac|lakh)?\b", re.I)
KEY_VALUE_RE = re.compile(r"^\s*([A-Za-z][A-Za-z0-9 /&()._-]{1,60})\s*[:=-]\s*(.{1,300})\s*$")
ID_RE = re.compile(r"\b[A-Z]{1,5}-\d{2,}\b|\b[A-Z]{2,}\d{4,}\b")
AGREEMENT_ROW_ID_RE = re.compile(r"^[RC]A-\d{3,}$", re.I)
PROPERTY_ROW_ID_RE = re.compile(r"^P[RC]-\d{3,}$", re.I)
BULLET_RE = re.compile(r"^\s*(?:[-*•]|\(\w+\)|[a-zA-Z]\)|\d+[.)])\s+(.+)$")
QUESTION_SIGNAL_RE = re.compile(r"\b(answer\s+key|choose\s+the\s+correct|multiple\s+choice|mcq|options?|questions?)\b", re.I)
REPORT_SIGNAL_RE = re.compile(r"\b(report|executive summary|findings|recommendations|analysis|overview)\b", re.I)
NOTICE_SIGNAL_RE = re.compile(r"\b(notice|subject|to whom it may concern|hereby notified|effective date)\b", re.I)
LETTER_SIGNAL_RE = re.compile(r"\b(dear\s+sir|dear\s+madam|dear\s+[a-z]+|sincerely|yours faithfully|regards)\b", re.I)
INVOICE_SIGNAL_RE = re.compile(r"\b(invoice|receipt|bill to|subtotal|grand total|tax|gst|amount due)\b", re.I)
CONTRACT_SIGNAL_RE = re.compile(r"\b(agreement|contract|deed|party|parties|whereas|terms and conditions|clause|article)\b", re.I)


def _tenant_dir(tenant_id: str | None) -> Path:
    settings = load_settings()
    path = settings.structured_dir / safe_workspace_id(tenant_id)
    path.mkdir(parents=True, exist_ok=True)
    return path


def _index_path(tenant_id: str | None) -> Path:
    return _tenant_dir(tenant_id) / "documents.json"


def _read_index(tenant_id: str | None) -> dict[str, Any]:
    path = _index_path(tenant_id)
    if not path.exists():
        return {"version": 1, "documents": []}
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except json.JSONDecodeError:
        return {"version": 1, "documents": []}


def _write_index(tenant_id: str | None, payload: dict[str, Any]) -> None:
    _index_path(tenant_id).write_text(json.dumps(payload, indent=2, ensure_ascii=False), encoding="utf-8")


def structured_documents(tenant_id: str | None) -> list[dict[str, Any]]:
    return list(_read_index(tenant_id).get("documents", []))


def remove_structured_document(path: Path, tenant_id: str | None) -> bool:
    resolved = str(path.resolve()).lower()
    index = _read_index(tenant_id)
    documents = index.get("documents", [])
    kept = [document for document in documents if document.get("document_id") != resolved]
    if len(kept) == len(documents):
        return False
    index["documents"] = kept
    _write_index(tenant_id, index)
    return True


def clear_structured_documents(tenant_id: str | None) -> int:
    documents = structured_documents(tenant_id)
    _write_index(tenant_id, {"version": 1, "documents": []})
    return len(documents)


def _money(value: str) -> int:
    return int(re.sub(r"[^\d]", "", value) or "0")


def _parse_money(value: str) -> int:
    text = value.lower()
    match = re.search(r"\d[\d,]*(?:\.\d+)?", text)
    digits = match.group(0).replace(",", "") if match else ""
    if not digits:
        return 0
    amount = float(digits)
    if re.search(r"\b(cr|crore)\b", text):
        amount *= 10_000_000
    elif re.search(r"\b(l|lac|lakh)\b", text):
        amount *= 100_000
    return int(round(amount))


def _parse_date(value: str) -> date | None:
    value = value.strip()
    for fmt in ("%d-%b-%y", "%d-%b-%Y", "%d/%b/%y", "%d/%b/%Y", "%d-%m-%y", "%d-%m-%Y", "%d/%m/%y", "%d/%m/%Y"):
        try:
            parsed = datetime.strptime(value, fmt).date()
        except ValueError:
            continue
        if parsed.year < 100:
            parsed = parsed.replace(year=parsed.year + 2000)
        return parsed
    return None


def _date_payload(value: str, prefix: str) -> dict[str, Any]:
    parsed = _parse_date(value)
    if not parsed:
        return {prefix: value}
    return {prefix: value, f"{prefix}_iso": parsed.isoformat()}


def _find_line_index(lines: list[str], marker: str, start: int = 0) -> int | None:
    marker_lower = marker.lower()
    for index in range(start, len(lines)):
        if lines[index].lower().startswith(marker_lower):
            return index
    return None


def _text_for_file(path: Path, tenant_id: str | None = None) -> str:
    suffix = path.suffix.lower()
    if suffix in CONVERTIBLE_EXTENSIONS and not is_markdown_sidecar(path):
        text = readable_text(path, tenant_id)
        if text.strip():
            return text
    if is_image_file(path):
        return extract_image_text(path)
    if suffix in {".txt", ".md"}:
        return path.read_text(encoding="utf-8", errors="ignore")
    if suffix == ".pdf":
        reader = PdfReader(path)
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    if suffix == ".docx":
        from docx import Document

        document = Document(path)
        parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if values:
                    parts.append(" | ".join(values))
        return "\n".join(parts)
    if suffix == ".xlsx":
        workbook = load_workbook(path, data_only=True)
        parts: list[str] = []
        for sheet in workbook.worksheets:
            parts.append(f"Sheet: {sheet.title}")
            for row in sheet.iter_rows(values_only=True):
                values = [str(value).strip() for value in row if value not in (None, "")]
                if values:
                    parts.append(" | ".join(values))
        return "\n".join(parts)
    return ""


def _lines(text: str) -> list[str]:
    return [line.strip() for line in _strip_conversion_metadata(text).splitlines() if line.strip()]


def _strip_conversion_metadata(text: str) -> str:
    return re.sub(r"^\s*<!--\s*AutoNova Markdown conversion.*?-->\s*", "", text, flags=re.S)


def _classify_document(path: Path, text: str) -> str:
    if path.suffix.lower() == ".xlsx":
        return "spreadsheet"
    haystack = f"{path.name}\n{text[:1200]}".lower()
    if "property" in haystack and ("listing" in haystack or "portfolio" in haystack):
        return "property_listing"
    if "rent" in haystack and ("collection" in haystack or "tracker" in haystack):
        return "rent_collection_tracker"
    if "agreement" in haystack and ("tenant" in haystack or "landlord" in haystack or "rent" in haystack):
        return "rental_agreement"
    if ("client" in haystack or "contact" in haystack) and ("phone" in haystack or "email" in haystack):
        return "client_contacts"
    if _looks_like_question_bank(text):
        return "question_bank"
    if INVOICE_SIGNAL_RE.search(haystack):
        return "invoice_receipt"
    if NOTICE_SIGNAL_RE.search(haystack):
        return "notice"
    if LETTER_SIGNAL_RE.search(haystack):
        return "letter"
    if REPORT_SIGNAL_RE.search(haystack):
        return "report"
    if CONTRACT_SIGNAL_RE.search(haystack):
        return "contract_legal"
    return "generic_document"


def _looks_like_question_bank(text: str) -> bool:
    haystack = text[:2500]
    if not QUESTION_SIGNAL_RE.search(haystack) and "coding" not in haystack.lower():
        return False
    question_marks = haystack.count("?")
    options = len(MCQ_OPTION_RE.findall(haystack))
    numbered_questions = len(QUESTION_TITLE_RE.findall(haystack))
    answer_keys = len(ANSWER_KEY_RE.findall(haystack))
    return question_marks >= 2 or options >= 4 or answer_keys >= 2 or numbered_questions >= 3


def _parse_xlsx_records(path: Path) -> list[dict[str, Any]]:
    workbook = load_workbook(path, data_only=True)
    records: list[dict[str, Any]] = []
    for sheet in workbook.worksheets:
        rows = [
            ["" if value is None else str(value).strip() for value in row]
            for row in sheet.iter_rows(values_only=True)
        ]
        rows = [row for row in rows if any(cell for cell in row)]
        if not rows:
            continue
        header_index = _header_row_index(rows)
        headers = _clean_headers(rows[header_index]) if header_index is not None else []
        if headers:
            for row_number, row in enumerate(rows[header_index + 1 :], start=header_index + 2):
                values = row[: len(headers)]
                if not any(values):
                    continue
                fields = {
                    header: values[index]
                    for index, header in enumerate(headers)
                    if index < len(values) and values[index] != ""
                }
                if fields:
                    records.append(
                        {
                            "kind": "table_row",
                            "sheet": sheet.title,
                            "row_number": row_number,
                            "fields": fields,
                            "source_file": str(path),
                        }
                    )
        else:
            for row_number, row in enumerate(rows, start=1):
                records.append(
                    {
                        "kind": "spreadsheet_row",
                        "sheet": sheet.title,
                        "row_number": row_number,
                        "values": [value for value in row if value],
                        "source_file": str(path),
                    }
                )
    return records


def _parse_docx_table_records(path: Path) -> list[dict[str, Any]]:
    from docx import Document

    document = Document(path)
    records: list[dict[str, Any]] = []
    for table_index, table in enumerate(document.tables, start=1):
        rows = [
            [cell.text.strip() for cell in row.cells]
            for row in table.rows
            if any(cell.text.strip() for cell in row.cells)
        ]
        if not rows:
            continue
        header_index = _header_row_index(rows)
        headers = _clean_headers(rows[header_index]) if header_index is not None else []
        if headers:
            for row_number, row in enumerate(rows[header_index + 1 :], start=header_index + 2):
                values = row[: len(headers)]
                fields = {
                    header: values[index]
                    for index, header in enumerate(headers)
                    if index < len(values) and values[index] != ""
                }
                if fields:
                    records.append(
                        {
                            "kind": "table_row",
                            "table": table_index,
                            "row_number": row_number,
                            "fields": fields,
                            "source_file": str(path),
                        }
                    )
        else:
            for row_number, row in enumerate(rows, start=1):
                records.append(
                    {
                        "kind": "document_table_row",
                        "table": table_index,
                        "row_number": row_number,
                        "values": [value for value in row if value],
                        "source_file": str(path),
                    }
                )
    return records


def _header_row_index(rows: list[list[str]]) -> int | None:
    for index, row in enumerate(rows[:10]):
        non_empty = [cell for cell in row if cell]
        if len(non_empty) >= 2 and sum(1 for cell in non_empty if re.search(r"[A-Za-z]", cell)) >= 2:
            return index
    return None


def _clean_headers(row: list[str]) -> list[str]:
    headers: list[str] = []
    seen: Counter[str] = Counter()
    for index, value in enumerate(row):
        header = re.sub(r"[^a-zA-Z0-9]+", "_", value.strip().lower()).strip("_")
        header = header or f"column_{index + 1}"
        seen[header] += 1
        if seen[header] > 1:
            header = f"{header}_{seen[header]}"
        headers.append(header)
    return headers


def _parse_universal_records(text: str, source: Path, document_type: str) -> list[dict[str, Any]]:
    lines = _lines(text)
    records: list[dict[str, Any]] = [_document_profile_record(lines, source, document_type)]
    records.extend(_parse_key_value_records(lines, source))
    records.extend(_parse_entity_records(lines, source))
    records.extend(_parse_outline_records(lines, source))
    records.extend(_parse_paragraph_records(text, source))
    records.extend(_parse_bullet_records(lines, source))
    records.extend(_parse_inline_table_records(lines, source))
    records.extend(_parse_stacked_table_records(lines, source))
    records.extend(_parse_qa_pair_records(lines, source))
    return _dedupe_universal_records(records)[:800]


def _document_profile_record(lines: list[str], source: Path, document_type: str) -> dict[str, Any]:
    title = next((line.lstrip("# ").strip() for line in lines[:8] if len(line.lstrip("# ").strip()) > 3), source.stem)
    return {
        "kind": "document_profile",
        "title": title[:180],
        "document_type": document_type,
        "line_count": len(lines),
        "source_file": str(source),
    }


def _parse_key_value_records(lines: list[str], source: Path) -> list[dict[str, Any]]:
    records: list[dict[str, Any]] = []
    for index, line in enumerate(lines):
        if _split_table_line(line):
            continue
        key_value = KEY_VALUE_RE.match(line)
        if not key_value:
            continue
        key = key_value.group(1).strip()
        value = key_value.group(2).strip()
        records.append(
            {
                "kind": "key_value",
                "key": key,
                "normalized_key": _normalize_key(key),
                "value": value,
                "normalized": _normalize_scalar(value),
                "source_line": index + 1,
                "source_file": str(source),
            }
        )
    return records


def _parse_entity_records(lines: list[str], source: Path) -> list[dict[str, Any]]:
    records: list[dict[str, Any]] = []
    for index, line in enumerate(lines):
        if _split_table_line(line):
            continue
        entities = _entities_from_line(line)
        if entities:
            records.append(
                {
                    "kind": "entity_line",
                    "text": line,
                    "entities": entities,
                    "source_line": index + 1,
                    "source_file": str(source),
                }
            )
    return records


def _parse_outline_records(lines: list[str], source: Path) -> list[dict[str, Any]]:
    records: list[dict[str, Any]] = []
    for index, line in enumerate(lines):
        if _looks_like_heading(line):
            records.append(
                {
                    "kind": "heading",
                    "title": line.rstrip(":"),
                    "level": _heading_level(line),
                    "source_line": index + 1,
                    "source_file": str(source),
                }
            )
    records.extend(_parse_section_records(lines, source))
    return records


def _parse_paragraph_records(text: str, source: Path) -> list[dict[str, Any]]:
    chunks = [re.sub(r"\s+", " ", part).strip() for part in re.split(r"\n\s*\n", text) if part.strip()]
    if len(chunks) <= 1:
        lines = _lines(text)
        chunks = []
        current: list[str] = []
        for line in lines:
            if _looks_like_heading(line) or KEY_VALUE_RE.match(line) or BULLET_RE.match(line):
                if current:
                    chunks.append(" ".join(current))
                    current = []
                continue
            current.append(line)
            if len(" ".join(current)) > 700:
                chunks.append(" ".join(current))
                current = []
        if current:
            chunks.append(" ".join(current))

    records: list[dict[str, Any]] = []
    for index, paragraph in enumerate(chunks, start=1):
        paragraph = paragraph.strip()
        if len(paragraph) < 40:
            continue
        records.append(
            {
                "kind": "paragraph",
                "paragraph_number": index,
                "text": paragraph[:1800],
                "word_count": len(paragraph.split()),
                "source_file": str(source),
            }
        )
    return records[:120]


def _parse_bullet_records(lines: list[str], source: Path) -> list[dict[str, Any]]:
    records: list[dict[str, Any]] = []
    for index, line in enumerate(lines):
        match = BULLET_RE.match(line)
        if not match:
            continue
        records.append(
            {
                "kind": "list_item",
                "text": match.group(1).strip(),
                "source_line": index + 1,
                "source_file": str(source),
            }
        )
    return records[:200]


def _parse_inline_table_records(lines: list[str], source: Path) -> list[dict[str, Any]]:
    records: list[dict[str, Any]] = []
    table_index = 0
    index = 0
    while index < len(lines):
        row = _split_table_line(lines[index])
        if len(row) < 2 or _is_markdown_table_separator(row):
            index += 1
            continue
        block: list[tuple[int, list[str]]] = []
        while index < len(lines):
            cells = _split_table_line(lines[index])
            if len(cells) < 2:
                break
            if not _is_markdown_table_separator(cells):
                block.append((index + 1, cells))
            index += 1
        if len(block) >= 2:
            header_index = _header_row_index([cells for _line, cells in block])
            table_index += 1
            table_title = _nearest_heading(lines, block[0][0] - 2)
            records.append(
                {
                    "kind": "table",
                    "table": table_index,
                    "title": table_title,
                    "row_count": max(len(block) - (1 if header_index is not None else 0), 0),
                    "source_line": block[0][0],
                    "source_file": str(source),
                }
            )
            if header_index is not None:
                headers = _clean_headers(block[header_index][1])
                for row_number, (source_line, cells) in enumerate(block[header_index + 1 :], start=1):
                    fields = {
                        header: _normalize_scalar(cells[position])
                        for position, header in enumerate(headers)
                        if position < len(cells) and cells[position] != ""
                    }
                    if fields:
                        records.append(
                            {
                                "kind": "table_row",
                                "table": table_index,
                                "table_title": table_title,
                                "row_number": row_number,
                                "fields": fields,
                                "source_line": source_line,
                                "source_file": str(source),
                            }
                        )
            else:
                for row_number, (source_line, cells) in enumerate(block, start=1):
                    records.append(
                        {
                            "kind": "table_row",
                            "table": table_index,
                            "table_title": table_title,
                            "row_number": row_number,
                            "values": [_normalize_scalar(cell) for cell in cells],
                            "source_line": source_line,
                            "source_file": str(source),
                        }
                    )
        index += 1
    return records[:300]


def _is_markdown_table_separator(cells: list[str]) -> bool:
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in cells)


def _nearest_heading(lines: list[str], before_index: int) -> str:
    for line in reversed(lines[: max(before_index + 1, 0)]):
        if line.startswith("#"):
            return line.lstrip("# ").strip()
        if _looks_like_heading(line) and not _split_table_line(line):
            return line.rstrip(":")
    return ""


def _parse_stacked_table_records(lines: list[str], source: Path) -> list[dict[str, Any]]:
    records: list[dict[str, Any]] = []
    table_index = 1000
    index = 0
    while index < len(lines):
        match = _stacked_table_at(lines, index)
        if not match:
            index += 1
            continue
        width, row_count = match
        headers = _clean_headers(lines[index : index + width])
        table_index += 1
        records.append(
            {
                "kind": "table",
                "table": table_index,
                "layout": "stacked_pdf_text",
                "row_count": row_count,
                "source_line": index + 1,
                "source_file": str(source),
            }
        )
        data_start = index + width
        for row_number in range(row_count):
            cells = lines[data_start + (row_number * width) : data_start + ((row_number + 1) * width)]
            fields = {
                header: _normalize_scalar(cells[position])
                for position, header in enumerate(headers)
                if position < len(cells) and cells[position] != ""
            }
            records.append(
                {
                    "kind": "table_row",
                    "table": table_index,
                    "layout": "stacked_pdf_text",
                    "row_number": row_number + 1,
                    "fields": fields,
                    "source_line": data_start + (row_number * width) + 1,
                    "source_file": str(source),
                }
            )
        index = data_start + (row_count * width)
    return records[:300]


def _stacked_table_at(lines: list[str], index: int) -> tuple[int, int] | None:
    for width in range(8, 1, -1):
        if index + (width * 3) > len(lines):
            continue
        headers = lines[index : index + width]
        if not _looks_like_stacked_headers(headers):
            continue
        data = lines[index + width :]
        row_count = 0
        for row_start in range(0, min(len(data), width * 25), width):
            row = data[row_start : row_start + width]
            if len(row) < width or not _looks_like_stacked_row(row):
                break
            row_count += 1
        if row_count >= 2:
            return width, row_count
    return None


def _looks_like_stacked_headers(headers: list[str]) -> bool:
    if len(headers) < 2:
        return False
    joined = " ".join(headers).lower()
    first_header = _normalize_key(headers[0])
    known_first_headers = {
        "id",
        "sr_no",
        "serial",
        "no",
        "agr_id",
        "agreement_id",
        "invoice",
        "date",
        "name",
        "tenant",
        "description",
        "item",
        "particulars",
    }
    if first_header not in known_first_headers and not any(first_header.endswith(f"_{key}") for key in known_first_headers):
        return False
    if len(set(_normalize_key(header) for header in headers)) < len(headers):
        return False
    if not any(keyword in joined for keyword in ("date", "name", "id", "amount", "rent", "status", "property", "tenant", "total", "description", "price", "qty")):
        return False
    for header in headers:
        if len(header) > 45 or not re.search(r"[A-Za-z]", header):
            return False
        if DATE_RE.fullmatch(header) or AMOUNT_RE.fullmatch(header) or ID_RE.fullmatch(header):
            return False
    return True


def _looks_like_stacked_row(row: list[str]) -> bool:
    joined = " ".join(row)
    if not (DATE_RE.search(row[0]) or AMOUNT_RE.search(row[0]) or ID_RE.search(row[0]) or re.search(r"\d", row[0])):
        return False
    signal_count = 0
    for cell in row:
        if DATE_RE.search(cell) or AMOUNT_RE.search(cell) or ID_RE.search(cell) or EMAIL_RE.search(cell) or PHONE_RE.search(cell):
            signal_count += 1
    return signal_count >= 1 and len(joined) <= 900


def _parse_qa_pair_records(lines: list[str], source: Path) -> list[dict[str, Any]]:
    records: list[dict[str, Any]] = []
    for index, line in enumerate(lines):
        if not re.search(r"\?$", line):
            continue
        answer_lines: list[str] = []
        for candidate in lines[index + 1 : index + 6]:
            if re.search(r"\?$", candidate) or _looks_like_heading(candidate):
                break
            answer_lines.append(candidate)
        if answer_lines:
            records.append(
                {
                    "kind": "qa_pair",
                    "question": line,
                    "answer": " ".join(answer_lines)[:1500],
                    "source_line": index + 1,
                    "source_file": str(source),
                }
            )
    return records[:100]


def _split_table_line(line: str) -> list[str]:
    if "|" in line:
        return [cell.strip() for cell in line.split("|") if cell.strip()]
    if "\t" in line:
        return [cell.strip() for cell in line.split("\t") if cell.strip()]
    if re.search(r"\S\s{2,}\S", line):
        return [cell.strip() for cell in re.split(r"\s{2,}", line) if cell.strip()]
    return []


def _normalize_key(value: str) -> str:
    return re.sub(r"[^a-z0-9]+", "_", value.lower()).strip("_")


def _normalize_scalar(value: Any) -> Any:
    if not isinstance(value, str):
        return value
    text = value.strip()
    parsed_date = _parse_date(text)
    if parsed_date:
        return {"raw": text, "date": parsed_date.isoformat()}
    if AMOUNT_RE.fullmatch(text) or re.fullmatch(r"(?:rs\.?|inr)?\s*\d[\d,]*(?:\.\d+)?\s*(?:cr|crore|l|lac|lakh)?", text, re.I):
        amount = _parse_money(text)
        if amount:
            return {"raw": text, "amount": amount}
    return text


def _heading_level(line: str) -> int:
    if re.match(r"^\d+\.\d+", line):
        return 3
    if re.match(r"^\d+[.)]", line):
        return 2
    return 1


def _dedupe_universal_records(records: list[dict[str, Any]]) -> list[dict[str, Any]]:
    seen: set[str] = set()
    deduped: list[dict[str, Any]] = []
    for record in records:
        key = json.dumps(
            {item_key: item_value for item_key, item_value in record.items() if item_key != "source_file"},
            sort_keys=True,
            ensure_ascii=False,
        )
        if key in seen:
            continue
        seen.add(key)
        deduped.append(record)
    return deduped


def _parse_generic_records(text: str, source: Path) -> list[dict[str, Any]]:
    records: list[dict[str, Any]] = []
    lines = _lines(text)
    for index, line in enumerate(lines):
        key_value = KEY_VALUE_RE.match(line)
        if key_value:
            records.append(
                {
                    "kind": "key_value",
                    "key": key_value.group(1).strip(),
                    "value": key_value.group(2).strip(),
                    "source_line": index + 1,
                    "source_file": str(source),
                }
            )

        entities = _entities_from_line(line)
        if entities:
            records.append(
                {
                    "kind": "entity_line",
                    "text": line,
                    "entities": entities,
                    "source_line": index + 1,
                    "source_file": str(source),
                }
            )

    records.extend(_parse_section_records(lines, source))
    return records[:300]


def _entities_from_line(line: str) -> dict[str, list[str]]:
    entities: dict[str, list[str]] = {}
    for label, pattern in (
        ("emails", EMAIL_RE),
        ("phones", PHONE_RE),
        ("dates", DATE_RE),
        ("amounts", AMOUNT_RE),
        ("ids", ID_RE),
    ):
        matches = [match.group(0).strip() for match in pattern.finditer(line)]
        if matches:
            entities[label] = list(dict.fromkeys(matches))
    if entities.get("dates"):
        normalized_dates = []
        for value in entities["dates"]:
            parsed = _parse_date(value)
            if parsed:
                normalized_dates.append(parsed.isoformat())
        if normalized_dates:
            entities["dates_iso"] = list(dict.fromkeys(normalized_dates))
    if entities.get("amounts"):
        normalized_amounts = []
        for value in entities["amounts"]:
            amount = _parse_money(value)
            if amount:
                normalized_amounts.append(str(amount))
        if normalized_amounts:
            entities["amounts_numeric"] = list(dict.fromkeys(normalized_amounts))
    return entities


def _parse_section_records(lines: list[str], source: Path) -> list[dict[str, Any]]:
    records: list[dict[str, Any]] = []
    current_title = ""
    current_lines: list[str] = []
    for line in lines:
        if _looks_like_heading(line):
            if current_title and current_lines:
                records.append(
                    {
                        "kind": "section",
                        "title": current_title,
                        "text": "\n".join(current_lines)[:2500],
                        "source_file": str(source),
                    }
                )
            current_title = line.rstrip(":")
            current_lines = []
        elif current_title:
            current_lines.append(line)
    if current_title and current_lines:
        records.append(
            {
                "kind": "section",
                "title": current_title,
                "text": "\n".join(current_lines)[:2500],
                "source_file": str(source),
            }
        )
    return records[:80]


def _looks_like_heading(line: str) -> bool:
    if len(line) > 90 or len(line.split()) > 12:
        return False
    if line.endswith(":"):
        return True
    return bool(re.match(r"^(\d+[\.)]\s+)?[A-Z][A-Za-z0-9 /&()._-]{3,}$", line)) and not any(
        pattern.search(line) for pattern in (EMAIL_RE, PHONE_RE, AMOUNT_RE)
    )


def _parse_property_records(text: str, source: Path) -> list[dict[str, Any]]:
    lines = _lines(text)
    records: list[dict[str, Any]] = []
    for index, line in enumerate(lines):
        match = PROPERTY_ID_RE.match(line)
        if not match:
            continue
        prop_id = match.group(1).upper()
        window = lines[index + 1 : index + 12]
        status = next((candidate for candidate in STATUSES if candidate.lower() in " ".join(window).lower()), "")
        records.append(
            {
                "kind": "property",
                "property_id": prop_id,
                "category": "residential" if prop_id.startswith("PR-") else "commercial",
                "location": window[0] if window else "",
                "status": status,
                "source_file": str(source),
                "source_line": index + 1,
            }
        )
    return records


def _parse_rent_records(text: str, source: Path) -> list[dict[str, Any]]:
    lines = _lines(text)
    records: list[dict[str, Any]] = []
    for index, line in enumerate(lines):
        match = AGREEMENT_ID_RE.match(line)
        if not match:
            continue
        agreement_id = match.group(1).upper()
        window = lines[index + 1 : index + 10]
        if len(window) < 7:
            continue
        rent_amount = _money(window[3])
        paid_amount = _money(window[4])
        status = window[7] if agreement_id.startswith("RA-") and len(window) > 7 else window[6]
        records.append(
            {
                "kind": "rent",
                "agreement_id": agreement_id,
                "tenant": window[0],
                "property_id": window[1],
                "rent_amount": rent_amount,
                "paid_amount": paid_amount,
                "pending_amount": max(rent_amount - paid_amount, 0),
                "status": status,
                "source_file": str(source),
                "source_line": index + 1,
            }
        )
    return records


def _value_text(value: Any) -> str:
    if isinstance(value, dict):
        return str(value.get("raw") or value.get("date") or value.get("amount") or "")
    return "" if value is None else str(value)


def _field_text(fields: dict[str, Any], *names: str) -> str:
    for name in names:
        value = fields.get(name)
        if value not in (None, ""):
            return _value_text(value).strip()
    return ""


def _field_amount(fields: dict[str, Any], *names: str) -> int:
    for name in names:
        value = fields.get(name)
        if isinstance(value, dict) and isinstance(value.get("amount"), int):
            return int(value["amount"])
        text = _value_text(value)
        if text:
            amount = _parse_money(text)
            if amount:
                return amount
    return 0


def _records_from_markdown_tables(records: list[dict[str, Any]], document_type: str, source: Path) -> list[dict[str, Any]]:
    if document_type == "property_listing":
        return _property_records_from_table_rows(records, source)
    if document_type == "rent_collection_tracker":
        return _rent_records_from_table_rows(records, source)
    return []


def _property_records_from_table_rows(records: list[dict[str, Any]], source: Path) -> list[dict[str, Any]]:
    parsed: list[dict[str, Any]] = []
    for record in records:
        if record.get("kind") != "table_row" or not isinstance(record.get("fields"), dict):
            continue
        fields = record["fields"]
        prop_id = _field_text(fields, "prop_id", "property_id", "id").upper()
        if not PROPERTY_ROW_ID_RE.match(prop_id):
            continue
        price_text = _field_text(fields, "price_inr", "price_rent_inr", "price", "rent")
        parsed.append(
            {
                "kind": "property",
                "property_id": prop_id,
                "category": "residential" if prop_id.startswith("PR-") else "commercial",
                "location": _field_text(fields, "location"),
                "property_type": _field_text(fields, "type"),
                "area_sqft": _field_amount(fields, "area_sqft", "area"),
                "floor": _field_text(fields, "floor"),
                "price": price_text,
                "price_amount": _parse_money(price_text),
                "status": _field_text(fields, "status"),
                "source_file": str(source),
                "source_line": record.get("source_line"),
            }
        )
    return parsed


def _rent_records_from_table_rows(records: list[dict[str, Any]], source: Path) -> list[dict[str, Any]]:
    parsed: list[dict[str, Any]] = []
    for record in records:
        if record.get("kind") != "table_row" or not isinstance(record.get("fields"), dict):
            continue
        fields = record["fields"]
        agreement_id = _field_text(fields, "agr_id", "agreement_id", "id").upper()
        if not AGREEMENT_ROW_ID_RE.match(agreement_id):
            continue
        rent_amount = _field_amount(fields, "rent_amt", "rent_amount", "rent")
        paid_amount = _field_amount(fields, "paid_amt", "paid_amount", "paid")
        parsed.append(
            {
                "kind": "rent",
                "agreement_id": agreement_id,
                "tenant": _field_text(fields, "tenant", "tenant_company", "company"),
                "property_id": _field_text(fields, "property", "property_id").upper(),
                "due_date": _field_text(fields, "due_date"),
                "paid_on": _field_text(fields, "paid_on"),
                "mode": _field_text(fields, "mode"),
                "rent_amount": rent_amount,
                "paid_amount": paid_amount,
                "pending_amount": max(rent_amount - paid_amount, 0),
                "status": _field_text(fields, "status"),
                "source_file": str(source),
                "source_line": record.get("source_line"),
            }
        )
    return parsed


def _parse_rental_agreement_records(text: str, source: Path) -> list[dict[str, Any]]:
    lines = _lines(text)
    records: list[dict[str, Any]] = []

    residential_start = _find_line_index(lines, "Active Residential Tenancies")
    commercial_start = _find_line_index(lines, "Active Commercial Tenancies")
    renewal_start = _find_line_index(lines, "Renewal Pipeline")

    if residential_start is not None:
        end = commercial_start if commercial_start is not None else len(lines)
        records.extend(_parse_agreement_table(lines[residential_start:end], source, "residential"))
    if commercial_start is not None:
        end = renewal_start if renewal_start is not None else len(lines)
        records.extend(_parse_agreement_table(lines[commercial_start:end], source, "commercial"))
    if renewal_start is not None:
        records.extend(_parse_renewal_pipeline_records(lines[renewal_start:], source))

    return records


def _parse_agreement_table(lines: list[str], source: Path, agreement_type: str) -> list[dict[str, Any]]:
    records: list[dict[str, Any]] = []
    ids = [
        index
        for index, line in enumerate(lines)
        if AGREEMENT_ROW_ID_RE.match(line) and index + 6 < len(lines)
    ]
    for index in ids:
        agreement_id = lines[index].upper()
        property_id = lines[index + 1].upper()
        if not PROPERTY_ROW_ID_RE.match(property_id):
            continue
        tenant = lines[index + 2]
        monthly_rent = _parse_money(lines[index + 3])
        record: dict[str, Any] = {
            "kind": "agreement",
            "agreement_id": agreement_id,
            "agreement_type": agreement_type,
            "property_id": property_id,
            "tenant": tenant,
            "monthly_rent": monthly_rent,
            "source_file": str(source),
        }
        if agreement_type == "residential":
            if index + 7 >= len(lines):
                continue
            record.update(
                {
                    "security_deposit": _parse_money(lines[index + 4]),
                    **_date_payload(lines[index + 5], "start_date"),
                    **_date_payload(lines[index + 6], "end_date"),
                    "status": lines[index + 7],
                }
            )
        else:
            record.update(
                {
                    **_date_payload(lines[index + 4], "start_date"),
                    **_date_payload(lines[index + 5], "end_date"),
                    "lock_in": lines[index + 6],
                    "status": lines[index + 7] if index + 7 < len(lines) else "",
                }
            )
        records.append(record)
    return records


def _parse_renewal_pipeline_records(lines: list[str], source: Path) -> list[dict[str, Any]]:
    records: list[dict[str, Any]] = []
    for index, line in enumerate(lines):
        if not AGREEMENT_ROW_ID_RE.match(line) or index + 5 >= len(lines):
            continue
        property_id = lines[index + 2].upper()
        expiry = _parse_date(lines[index + 3])
        if not PROPERTY_ROW_ID_RE.match(property_id) or not expiry:
            continue
        increase_match = re.search(r"\(([^)]+)\)", lines[index + 4])
        records.append(
            {
                "kind": "renewal_action",
                "agreement_id": line.upper(),
                "tenant": lines[index + 1],
                "property_id": property_id,
                "expiry": lines[index + 3],
                "expiry_iso": expiry.isoformat(),
                "proposed_new_rent": _parse_money(lines[index + 4]),
                "proposed_increase": increase_match.group(1) if increase_match else "",
                "action_owner": lines[index + 5],
                "source_file": str(source),
            }
        )
    return records


def _parse_client_contact_records(text: str, source: Path) -> list[dict[str, Any]]:
    lines = _lines(text)
    records: list[dict[str, Any]] = []
    sections = [
        (
            "tenant",
            "Active Tenants",
            ("Name", "Phone", "Email", "Property", "Type", "Since"),
            ("Prospective Buyers", "Vendors", "Service Partners"),
        ),
        (
            "buyer_lead",
            "Prospective Buyers",
            ("Name", "Phone", "Budget (INR)", "Requirement", "Last Contact", "Assigned To"),
            ("Vendors", "Service Partners"),
        ),
        (
            "vendor",
            "Vendors & Service Partners",
            ("Vendor Name", "Service", "Phone", "Rate", "Preferred"),
            ("Prestige Realty", "Internal Document", "Confidential"),
        ),
    ]
    for kind, start_marker, headers, end_markers in sections:
        section = _section_lines(lines, start_marker, end_markers)
        if not section:
            continue
        section = _drop_header_lines(section, headers)
        width = len(headers)
        for index in range(0, len(section), width):
            chunk = section[index : index + width]
            if len(chunk) < width:
                continue
            record = _contact_record(kind, headers, chunk, source)
            if record:
                records.append(record)
    return records


def _section_lines(lines: list[str], start_marker: str, end_markers: tuple[str, ...]) -> list[str]:
    start_index = next((index for index, line in enumerate(lines) if line.lower().startswith(start_marker.lower())), None)
    if start_index is None:
        return []
    end_index = len(lines)
    for index in range(start_index + 1, len(lines)):
        if any(marker.lower() in lines[index].lower() for marker in end_markers):
            end_index = index
            break
    return lines[start_index + 1 : end_index]


def _drop_header_lines(lines: list[str], headers: tuple[str, ...]) -> list[str]:
    header_set = {header.lower() for header in headers}
    return [line for line in lines if line.lower() not in header_set]


def _contact_record(kind: str, headers: tuple[str, ...], values: list[str], source: Path) -> dict[str, Any] | None:
    payload = dict(zip(headers, values))
    name = payload.get("Name") or payload.get("Vendor Name")
    phone = payload.get("Phone", "")
    email = payload.get("Email", "")
    if not name or PHONE_RE.search(name) or not PHONE_RE.search(phone):
        return None
    record: dict[str, Any] = {
        "kind": "contact",
        "contact_type": kind,
        "name": name,
        "phone": phone,
        "source_file": str(source),
    }
    if email and EMAIL_RE.search(email):
        record["email"] = email
    if kind == "tenant":
        record.update(
            {
                "property_id": payload.get("Property", ""),
                "property_type": payload.get("Type", ""),
                "since": payload.get("Since", ""),
            }
        )
    elif kind == "buyer_lead":
        record.update(
            {
                "budget": payload.get("Budget (INR)", ""),
                "requirement": payload.get("Requirement", ""),
                "last_contact": payload.get("Last Contact", ""),
                "assigned_to": payload.get("Assigned To", ""),
            }
        )
    elif kind == "vendor":
        record.update(
            {
                "service": payload.get("Service", ""),
                "rate": payload.get("Rate", ""),
                "preferred": payload.get("Preferred", ""),
            }
        )
    return record


def _parse_question_records(text: str, source: Path) -> list[dict[str, Any]]:
    code_records = _parse_code_question_records(text, source)
    if code_records:
        return code_records

    answer_key = _parse_answer_key(text)
    mcq_records = _parse_mcq_records(text, source, answer_key)
    if mcq_records:
        return mcq_records

    seen: set[str] = set()
    records: list[dict[str, Any]] = []
    for match in QUESTION_TITLE_RE.finditer(text):
        number = int(match.group(1))
        title = re.sub(r"\s+", " ", match.group(2)).strip(" :-")
        if re.fullmatch(r"[A-D](?:,\s*\d{1,3}-[A-D])*", title, flags=re.I):
            continue
        key = f"{number}:{title.lower()}"
        if key in seen:
            continue
        seen.add(key)
        answer_letter = answer_key.get(number)
        records.append(
            {
                "kind": "question",
                "number": number,
                "title": title,
                "answer": answer_letter,
                "source_file": str(source),
            }
        )
    return records


def _parse_code_question_records(text: str, source: Path) -> list[dict[str, Any]]:
    records: list[dict[str, Any]] = []
    pattern = re.compile(
        r"(?ms)^#{1,3}\s*(\d{1,3})\.\s+([^\n]+)\n(.*?)(?=^#{1,3}\s*\d{1,3}\.\s+|\Z)"
    )
    for match in pattern.finditer(_strip_conversion_metadata(text)):
        number = int(match.group(1))
        title = re.sub(r"\s+", " ", match.group(2)).strip()
        body = match.group(3).strip()
        code_blocks = re.findall(r"```(?:[A-Za-z0-9_+-]+)?\s*(.*?)```", body, flags=re.S)
        code = "\n".join(block.strip() for block in code_blocks if block.strip()).strip()
        if not title:
            continue
        records.append(
            {
                "kind": "question",
                "number": number,
                "title": title,
                "question_type": "coding",
                "answer_text": code or re.sub(r"\s+", " ", body)[:2500],
                "source_file": str(source),
            }
        )
    return records[:300]


def _parse_answer_key(text: str) -> dict[int, str]:
    marker = re.search(r"\banswer\s+key\b", text, flags=re.I)
    key_text = text[marker.end() :] if marker else text
    return {int(number): letter.upper() for number, letter in ANSWER_KEY_RE.findall(key_text)}


def _parse_mcq_records(text: str, source: Path, answer_key: dict[int, str]) -> list[dict[str, Any]]:
    lines = _lines(text)
    records: list[dict[str, Any]] = []
    current: dict[str, Any] | None = None
    for line in lines:
        if re.search(r"\banswer\s+key\b", line, flags=re.I):
            break
        question_match = MCQ_QUESTION_RE.match(line)
        option_match = MCQ_OPTION_RE.match(line)
        if question_match:
            if current:
                _finalize_mcq_record(current, answer_key)
                records.append(current)
            number = int(question_match.group(1))
            current = {
                "kind": "question",
                "number": number,
                "title": re.sub(r"\s+", " ", question_match.group(2)).strip(" :-"),
                "options": {},
                "source_file": str(source),
            }
        elif option_match and current:
            letter = option_match.group(1).upper()
            current["options"][letter] = re.sub(r"\s+", " ", option_match.group(2)).strip()
    if current:
        _finalize_mcq_record(current, answer_key)
        records.append(current)
    return records


def _finalize_mcq_record(record: dict[str, Any], answer_key: dict[int, str]) -> None:
    answer = answer_key.get(int(record["number"]))
    if answer:
        record["answer"] = answer
        record["answer_text"] = record.get("options", {}).get(answer, "")


def _parse_reference_records(text: str, source: Path) -> list[dict[str, Any]]:
    matches = list(REFERENCE_ID_RE.finditer(text))
    records: list[dict[str, Any]] = []
    for index, match in enumerate(matches):
        start = match.end()
        end = matches[index + 1].start() if index + 1 < len(matches) else len(text)
        name = re.sub(r"\s+", " ", text[start:end]).strip(" |,-:")
        name = re.sub(r"\b(reference|candidate|name)\b", "", name, flags=re.I).strip(" |,-:")
        if name:
            records.append(
                {
                    "kind": "reference_record",
                    "reference_id": match.group(0),
                    "name": name,
                    "source_file": str(source),
                }
            )
    return records


def _parse_clauses(text: str, source: Path) -> list[dict[str, Any]]:
    clauses: list[dict[str, Any]] = []
    current_title = ""
    current_lines: list[str] = []
    for line in _lines(text):
        is_heading = bool(re.match(r"^(\d+[\.)]\s+)?[A-Z][A-Za-z /&-]{3,80}:?$", line)) and len(line.split()) <= 10
        if is_heading:
            if current_title and current_lines:
                clauses.append({"title": current_title, "text": "\n".join(current_lines), "source_file": str(source)})
            current_title = line.rstrip(":")
            current_lines = []
        elif current_title:
            current_lines.append(line)
    if current_title and current_lines:
        clauses.append({"title": current_title, "text": "\n".join(current_lines), "source_file": str(source)})
    return clauses[:80]


def decode_file(path: Path, tenant_id: str | None) -> dict[str, Any] | None:
    if path.suffix.lower() not in SUPPORTED_TEXT_EXTENSIONS or is_markdown_sidecar(path):
        return None
    text = _text_for_file(path, tenant_id)
    if not text.strip():
        return None

    document_type = _classify_document(path, text)
    records: list[dict[str, Any]] = _parse_universal_records(text, path, document_type)
    records.extend(_records_from_markdown_tables(records, document_type, path))
    if path.suffix.lower() == ".xlsx":
        records.extend(_parse_xlsx_records(path))
    elif path.suffix.lower() == ".docx":
        records.extend(_parse_docx_table_records(path))
    if document_type == "property_listing":
        records.extend(_parse_property_records(text, path))
    elif document_type == "rent_collection_tracker":
        records.extend(_parse_rent_records(text, path))
    elif document_type == "client_contacts":
        records.extend(_parse_client_contact_records(text, path))
    elif document_type == "rental_agreement":
        records.extend(_parse_rental_agreement_records(text, path))
        records.extend(_parse_reference_records(text, path))
    elif document_type == "question_bank":
        records.extend(_parse_question_records(text, path))
        records.extend(_parse_reference_records(text, path))
    else:
        records.extend(_parse_reference_records(text, path))

    return {
        "document_id": str(path.resolve()).lower(),
        "source_file": str(path),
        "file_name": path.name,
        "document_type": document_type,
        "decoded_at": datetime.now().replace(microsecond=0).isoformat(),
        "summary": {
            "char_count": len(text),
            "line_count": len(_lines(text)),
            "record_count": len(records),
            "clause_count": len(_parse_clauses(text, path)) if document_type in {"rental_agreement", "generic_document"} else 0,
        },
        "records": records,
        "clauses": _parse_clauses(text, path) if document_type in {"rental_agreement", "generic_document"} else [],
    }


def upsert_document(path: Path, tenant_id: str | None) -> dict[str, Any] | None:
    decoded = decode_file(path, tenant_id)
    if not decoded:
        return None
    index = _read_index(tenant_id)
    documents = [
        document
        for document in index.get("documents", [])
        if document.get("document_id") != decoded["document_id"]
    ]
    documents.append(decoded)
    index["documents"] = sorted(documents, key=lambda item: item.get("decoded_at", ""), reverse=True)
    _write_index(tenant_id, index)
    return decoded


def structure_existing_files(tenant_id: str | None = None) -> int:
    settings = load_settings()
    workspace = safe_workspace_id(tenant_id)
    roots = [
        settings.uploads_dir / workspace if tenant_id else settings.uploads_dir,
        settings.generated_dir / workspace if tenant_id else settings.generated_dir,
        settings.data_dir / workspace if tenant_id else settings.data_dir,
    ]
    documents: list[dict[str, Any]] = []
    for root in roots:
        if not root.exists():
            continue
        paths = root.rglob("*") if tenant_id else root.glob("*")
        for path in paths:
            if path.is_file() and path.suffix.lower() in SUPPORTED_TEXT_EXTENSIONS and not is_markdown_sidecar(path):
                decoded = decode_file(path, tenant_id)
                if decoded:
                    documents.append(decoded)
    documents = sorted(documents, key=lambda item: item.get("decoded_at", ""), reverse=True)
    _write_index(tenant_id, {"version": 1, "documents": documents})
    return len(documents)


def structured_context(query: str, tenant_id: str | None, limit: int = 12) -> str:
    query_tokens = {token for token in re.findall(r"[a-z0-9]+", query.lower()) if len(token) > 2}
    if not query_tokens:
        return ""
    index = _read_index(tenant_id)
    matches: list[tuple[int, str]] = []
    for document in index.get("documents", []):
        doc_text = f"{document.get('file_name', '')} {document.get('document_type', '')}"
        for record in document.get("records", []):
            haystack = f"{doc_text} {json.dumps(record, ensure_ascii=False)}".lower()
            score = sum(1 for token in query_tokens if token in haystack)
            if score:
                matches.append((score, _format_context_record(record, document)))
        for clause in document.get("clauses", []):
            haystack = f"{doc_text} {clause.get('title', '')} {clause.get('text', '')}".lower()
            score = sum(1 for token in query_tokens if token in haystack)
            if score:
                text = clause.get("text", "")[:1200]
                matches.append((score, f"Clause: {clause.get('title')}\n{text}"))
    matches.sort(key=lambda item: item[0], reverse=True)
    return "\n\n".join(text for _score, text in matches[:limit])


def _format_context_record(record: dict[str, Any], document: dict[str, Any]) -> str:
    prefix = f"{document.get('file_name', 'Document')}: "
    kind = record.get("kind", "record")
    if kind == "agreement":
        return (
            f"{prefix}{record.get('agreement_id')} - {record.get('tenant')} ({record.get('property_id')}). "
            f"Rent {_money_label(int(record.get('monthly_rent') or 0))}, "
            f"start {record.get('start_date')}, end {record.get('end_date')}, status {record.get('status')}."
        )
    if kind == "renewal_action":
        return (
            f"{prefix}{record.get('agreement_id')} renewal - {record.get('tenant')} ({record.get('property_id')}). "
            f"Expiry {record.get('expiry')}, proposed rent {_money_label(int(record.get('proposed_new_rent') or 0))}, "
            f"owner {record.get('action_owner')}."
        )
    if kind == "contact":
        parts = [str(record.get("name", "")), str(record.get("contact_type", ""))]
        for field in ("phone", "email", "property_id", "requirement", "service", "assigned_to"):
            if record.get(field):
                parts.append(f"{field.replace('_', ' ')}: {record[field]}")
        return prefix + ", ".join(part for part in parts if part)
    if kind == "table_row":
        return prefix + _record_summary(record)
    if kind in {"paragraph", "section", "heading", "key_value", "qa_pair", "list_item"}:
        return prefix + _record_summary(record)
    return prefix + _record_summary(record)


def _dedupe_records(records: list[dict[str, Any]], key_fields: tuple[str, ...]) -> list[dict[str, Any]]:
    deduped: dict[tuple[str, ...], dict[str, Any]] = {}
    for record in records:
        key = tuple(str(record.get(field, "")).lower() for field in key_fields)
        if any(key):
            deduped.setdefault(key, record)
    return list(deduped.values())


def _tokens(text: str) -> set[str]:
    return {token for token in re.findall(r"[a-z0-9]+", text.lower()) if len(token) > 1}


def _answer_question_from_records(query: str, records: list[dict[str, Any]]) -> str | None:
    if not re.search(
        r"\b(question|questions|problem|problems|code|coding|program|answer|mcq|option|tcs|nqt|java|python|prime|factorial|fibonacci|array|string|matrix|sort|search|q\s*\d{1,3})\b",
        query,
        re.I,
    ):
        return None
    questions = _dedupe_records(
        [record for record in records if record.get("kind") == "question"],
        ("source_file", "number", "title"),
    )
    if not questions:
        return None

    number_match = re.search(r"\b(?:q|question|no\.?|number)?\s*(\d{1,3})\b", query, flags=re.I)
    selected: dict[str, Any] | None = None
    if number_match and re.search(r"\b(answer|correct|question|q|no\.?|number)\b", query, flags=re.I):
        number = int(number_match.group(1))
        matching_number = [record for record in questions if int(record.get("number") or 0) == number]
        selected = next((record for record in matching_number if record.get("answer")), None)
        selected = selected or (matching_number[0] if matching_number else None)

    if not selected:
        query_tokens = _tokens(query)
        scored: list[tuple[int, dict[str, Any]]] = []
        for record in questions:
            title_tokens = _tokens(str(record.get("title", "")))
            if not title_tokens:
                continue
            score = len(query_tokens & title_tokens)
            if score:
                scored.append((score, record))
        if scored:
            scored.sort(key=lambda item: item[0], reverse=True)
            best_score = scored[0][0]
            best_candidates = [record for score, record in scored if score == best_score]
            best_record = next((record for record in best_candidates if record.get("answer")), best_candidates[0])
            if best_score >= max(1, min(2, len(_tokens(str(best_record.get("title", "")))))):
                selected = best_record

    if not selected:
        return None
    if selected.get("answer_text") and not selected.get("answer"):
        title = str(selected.get("title", "")).strip()
        number = selected.get("number")
        heading = f"{number}. {title}" if number else title
        return f"{heading}\n{selected.get('answer_text')}".strip()
    if not selected.get("answer"):
        return None
    answer = str(selected.get("answer", "")).upper()
    answer_text = str(selected.get("answer_text", "")).strip()
    if answer_text:
        return f"{answer}) {answer_text}"
    return answer


def _answer_qa_pair_query(query: str, documents: list[dict[str, Any]], records: list[dict[str, Any]]) -> str | None:
    qa_sources = {
        str(document.get("source_file", ""))
        for document in documents
        if document.get("document_type") == "question_bank" or re.search(r"\b(faq|questions?)\b", str(document.get("file_name", "")), re.I)
    }
    qa_pairs = [
        record
        for record in records
        if record.get("kind") == "qa_pair" and (not qa_sources or str(record.get("source_file", "")) in qa_sources)
    ]
    if not qa_pairs or not re.search(r"\b(faq|question|answer|say|what|how|can|is|does|process|documents?|rent|deposit|maintenance)\b", query, re.I):
        return None

    qa_stopwords = {
        "what", "when", "where", "which", "who", "how", "why", "does", "say", "said", "according",
        "document", "documents", "uploaded", "file", "files", "faq", "question", "answer", "process",
        "list", "show", "give", "details", "detail", "about",
    }
    query_tokens = {token for token in _tokens(query) if token not in qa_stopwords and len(token) > 2}
    if not query_tokens:
        return None
    scored: list[tuple[int, dict[str, Any]]] = []
    for record in qa_pairs:
        haystack = f"{record.get('question', '')} {record.get('answer', '')}"
        score = len(query_tokens & _tokens(haystack))
        if score:
            scored.append((score, record))
    if not scored:
        return None
    scored.sort(key=lambda item: item[0], reverse=True)
    best_score, best = scored[0]
    if best_score < 2 and len(query_tokens) > 2:
        return None
    return f"{best.get('question')}\n{best.get('answer')}".strip()


def _answer_contact_query(query: str, records: list[dict[str, Any]]) -> str | None:
    contacts = _dedupe_records(
        [record for record in records if record.get("kind") == "contact"],
        ("contact_type", "name"),
    )
    if not contacts or not re.search(r"\b(contacts?|phone|email|number|leads?|tenants?|vendors?|buyers?|assigned|budget|requirement)\b", query, re.I):
        return None
    if re.search(r"\b(how many|count|number of|total)\b", query, re.I):
        counts = Counter(str(record.get("contact_type", "contact")) for record in contacts)
        parts = ", ".join(f"{kind}: {count}" for kind, count in sorted(counts.items()))
        return f"I found {len(contacts)} contact(s)." + (f" Breakdown: {parts}." if parts else "")
    query_tokens = _tokens(query)
    scored: list[tuple[int, dict[str, Any]]] = []
    for record in contacts:
        haystack = " ".join(str(value) for key, value in record.items() if key != "source_file")
        score = len(query_tokens & _tokens(haystack))
        if score:
            scored.append((score, record))
    if not scored:
        return None
    scored.sort(key=lambda item: item[0], reverse=True)
    selected = [record for score, record in scored if score == scored[0][0]][:5]
    lines = []
    for record in selected:
        detail_parts = [str(record.get("name"))]
        if record.get("contact_type"):
            detail_parts.append(f"type: {record['contact_type']}")
        if record.get("phone"):
            detail_parts.append(f"phone: {record['phone']}")
        if record.get("email"):
            detail_parts.append(f"email: {record['email']}")
        for field in ("property_id", "budget", "requirement", "assigned_to", "service", "rate", "preferred"):
            if record.get(field):
                detail_parts.append(f"{field.replace('_', ' ')}: {record[field]}")
        lines.append("- " + ", ".join(detail_parts))
    return "\n".join(lines)


def _expiry_window(query: str) -> int | None:
    match = re.search(r"\b(?:within|next|in)\s+(\d{1,4})\s*(days?|d|months?|mo|years?|yr)\b", query, re.I)
    if not match:
        return None
    value = int(match.group(1))
    unit = match.group(2).lower()
    if unit.startswith("month") or unit == "mo":
        return value * 30
    if unit.startswith("year") or unit == "yr":
        return value * 365
    return value


def _record_date(record: dict[str, Any], *fields: str) -> date | None:
    for field in fields:
        value = record.get(field)
        if not value:
            continue
        if isinstance(value, str):
            try:
                return date.fromisoformat(value)
            except ValueError:
                parsed = _parse_date(value)
                if parsed:
                    return parsed
    return None


def _reference_date_from_records(records: list[dict[str, Any]]) -> date:
    for record in records:
        text = " ".join(str(record.get(field, "")) for field in ("text", "title", "value"))
        month_match = re.search(
            r"\b(january|february|march|april|may|june|july|august|september|october|november|december)\s+(20\d{2})\b",
            text,
            re.I,
        )
        if month_match:
            try:
                return datetime.strptime(f"01 {month_match.group(1)} {month_match.group(2)}", "%d %B %Y").date()
            except ValueError:
                pass
    dated: list[date] = []
    for record in records:
        for value in record.get("entities", {}).get("dates_iso", []):
            try:
                dated.append(date.fromisoformat(value))
            except ValueError:
                continue
    return min(dated) if dated else date.today()


def _money_label(value: int) -> str:
    return f"Rs. {value:,}"


def _selected_agreements_for_query(query: str, agreements: list[dict[str, Any]]) -> list[dict[str, Any]]:
    lower = query.lower()
    selected = agreements
    if re.search(r"\bresidential\b", lower):
        selected = [record for record in selected if record.get("agreement_type") == "residential"]
    if re.search(r"\bcommercial\b", lower):
        selected = [record for record in selected if record.get("agreement_type") == "commercial"]
    for status in ("active", "expiring soon", "expired"):
        if status in lower:
            selected = [record for record in selected if str(record.get("status", "")).lower() == status]
    mentioned = [
        record
        for record in selected
        if any(str(record.get(field, "")).lower() in lower for field in ("agreement_id", "property_id") if record.get(field))
        or any(token in lower for token in re.findall(r"[a-z]{3,}", str(record.get("tenant", "")).lower()))
    ]
    return mentioned or selected


def _answer_agreement_query(query: str, records: list[dict[str, Any]]) -> str | None:
    agreements = _dedupe_records(
        [record for record in records if record.get("kind") == "agreement"],
        ("agreement_id",),
    )
    if not agreements or not re.search(r"\b(agreements?|leases?|tenants?|rent|expiry|expire|expires|expiring|deposit|property)\b", query, re.I):
        return None

    lower = query.lower()
    today = _reference_date_from_records(records)
    selected = _selected_agreements_for_query(query, agreements)

    window_days = _expiry_window(query)
    if window_days is not None or re.search(r"\b(expiry|expire|expires|expiring)\b", lower):
        expiry_selected = selected
        if window_days is not None:
            end = today + timedelta(days=window_days)
            expiry_selected = [
                record
                for record in selected
                if (expiry := _record_date(record, "end_date_iso", "end_date")) and today <= expiry <= end
            ]
        elif "expired" in lower:
            expiry_selected = [
                record
                for record in selected
                if (expiry := _record_date(record, "end_date_iso", "end_date")) and expiry < today
            ]
        elif "expiring soon" in lower:
            expiry_selected = [record for record in selected if str(record.get("status", "")).lower() == "expiring soon"]

        if re.search(r"\b(how many|count|number of|total)\b", lower):
            qualifier = f" within {window_days} day(s)" if window_days is not None else ""
            return f"{len(expiry_selected)} agreement(s) expire{qualifier} from {today.isoformat()}."

        if expiry_selected:
            lines = [f"{len(expiry_selected)} matching agreement(s):"]
            for record in sorted(expiry_selected, key=lambda item: item.get("end_date_iso", "")):
                lines.append(
                    f"- {record.get('agreement_id')} ({record.get('tenant')}, {record.get('property_id')}): "
                    f"ends {record.get('end_date')} / {record.get('end_date_iso')}, "
                    f"rent {_money_label(int(record.get('monthly_rent') or 0))}, status: {record.get('status')}."
                )
            return "\n".join(lines)
        return f"No matching agreement expiries were found from {today.isoformat()}."

    if re.search(r"\b(sum|total|average|avg|highest|lowest|max|min|monthly rent|deposit|security)\b", lower):
        field = "security_deposit" if re.search(r"\b(deposit|security)\b", lower) else "monthly_rent"
        values = [int(record.get(field) or 0) for record in selected if int(record.get(field) or 0) > 0]
        if not values:
            return None
        label = field.replace("_", " ")
        if re.search(r"\b(average|avg)\b", lower):
            return f"Average {label} is {_money_label(round(sum(values) / len(values)))} across {len(values)} agreement(s)."
        if re.search(r"\b(highest|max)\b", lower):
            record = max(selected, key=lambda item: int(item.get(field) or 0))
            return f"Highest {label} is {_money_label(int(record.get(field) or 0))}: {record.get('agreement_id')} ({record.get('tenant')})."
        if re.search(r"\b(lowest|min)\b", lower):
            record = min((item for item in selected if int(item.get(field) or 0) > 0), key=lambda item: int(item.get(field) or 0))
            return f"Lowest {label} is {_money_label(int(record.get(field) or 0))}: {record.get('agreement_id')} ({record.get('tenant')})."
        return f"Total {label} is {_money_label(sum(values))} across {len(values)} agreement(s)."

    if re.search(r"\b(how many|count|number of|total)\b", lower):
        counts = Counter(str(record.get("agreement_type", "agreement")) for record in selected)
        parts = ", ".join(f"{kind}: {count}" for kind, count in sorted(counts.items()))
        return f"I found {len(selected)} agreement(s)." + (f" Breakdown: {parts}." if parts else "")

    if re.search(r"\b(list|show|which|details?|who|what)\b", lower):
        visible = selected[:10]
        lines = [f"I found {len(selected)} matching agreement(s):"]
        for record in visible:
            details = _agreement_detail_parts(record)
            lines.append("- " + ", ".join(details))
        if len(selected) > len(visible):
            lines.append(f"{len(selected) - len(visible)} more matching agreement(s) are available.")
        return "\n".join(lines)

    return None


def _agreement_detail_parts(record: dict[str, Any]) -> list[str]:
    details = [
        str(record.get("agreement_id")),
        str(record.get("tenant")),
        str(record.get("property_id")),
        str(record.get("agreement_type")),
        f"rent {_money_label(int(record.get('monthly_rent') or 0))}",
    ]
    if record.get("security_deposit"):
        details.append(f"security deposit {_money_label(int(record.get('security_deposit') or 0))}")
    if record.get("start_date"):
        details.append(f"starts {record.get('start_date')}")
    if record.get("end_date"):
        details.append(f"ends {record.get('end_date')}")
    if record.get("lock_in"):
        details.append(f"lock-in {record.get('lock_in')}")
    if record.get("status"):
        details.append(f"status {record.get('status')}")
    return [part for part in details if part and part != "None"]


def _answer_renewal_query(query: str, records: list[dict[str, Any]]) -> str | None:
    renewals = _dedupe_records(
        [record for record in records if record.get("kind") == "renewal_action"],
        ("agreement_id",),
    )
    if not renewals or not re.search(r"\b(renewal|pipeline|action owner|proposed|new rent)\b", query, re.I):
        return None

    lower = query.lower()
    selected = renewals
    window_days = _expiry_window(query)
    if window_days is not None:
        today = date.today()
        end = today + timedelta(days=window_days)
        selected = [
            record
            for record in selected
            if (expiry := _record_date(record, "expiry_iso", "expiry")) and today <= expiry <= end
        ]

    if re.search(r"\b(how many|count|number of|total)\b", lower) and not re.search(r"\b(rent|amount|value)\b", lower):
        return f"I found {len(selected)} renewal action(s) in the structured renewal pipeline."

    if re.search(r"\b(sum|total|proposed|new rent|amount|value)\b", lower):
        total = sum(int(record.get("proposed_new_rent") or 0) for record in selected)
        return f"Total proposed new rent is {_money_label(total)} across {len(selected)} renewal action(s)."

    lines = [f"I found {len(selected)} renewal action(s):"]
    for record in selected[:10]:
        lines.append(
            f"- {record.get('agreement_id')} ({record.get('tenant')}, {record.get('property_id')}): "
            f"expires {record.get('expiry')} / {record.get('expiry_iso')}, proposed rent "
            f"{_money_label(int(record.get('proposed_new_rent') or 0))}, owner: {record.get('action_owner')}."
        )
    return "\n".join(lines)


def _flatten_numeric_fields(record: dict[str, Any]) -> dict[str, int]:
    values: dict[str, int] = {}
    fields = record.get("fields")
    if isinstance(fields, dict):
        for key, value in fields.items():
            if isinstance(value, dict) and isinstance(value.get("amount"), int):
                values[key] = int(value["amount"])
            elif isinstance(value, str):
                amount = _parse_money(value)
                if amount and re.search(r"\d", value):
                    values[key] = amount
    for key, value in record.items():
        if isinstance(value, int):
            values[key] = value
        elif isinstance(value, dict) and isinstance(value.get("amount"), int):
            values[key] = int(value["amount"])
    return values


def _answer_general_structured_query(query: str, documents: list[dict[str, Any]], records: list[dict[str, Any]]) -> str | None:
    lower = query.lower()
    countable_kinds = {
        "table_row": "rows?|records?|entries?",
        "table": "tables?",
        "paragraph": "paragraphs?",
        "heading": "headings?|sections?",
        "section": "sections?|clauses?",
        "list_item": "items?|points?|bullets?",
        "key_value": "fields?|key values?|metadata",
        "qa_pair": "questions?|answers?|qa|q&a",
    }

    if re.search(r"\b(how many|count|number of|total)\b", lower):
        for kind, pattern in countable_kinds.items():
            if re.search(rf"\b({pattern})\b", lower):
                selected = [record for record in records if record.get("kind") == kind]
                return f"I found {len(selected)} structured {kind.replace('_', ' ')} record(s)."

    if re.search(r"\b(list|show|give|find|which|details?)\b", lower):
        for kind, pattern in countable_kinds.items():
            if re.search(rf"\b({pattern})\b", lower):
                selected = [record for record in records if record.get("kind") == kind]
                if not selected:
                    return None
                lines = [f"I found {len(selected)} structured {kind.replace('_', ' ')} record(s)."]
                for index, record in enumerate(selected[:10], start=1):
                    lines.append(f"{index}. {_record_summary(record)}")
                if len(selected) > 10:
                    lines.append(f"{len(selected) - 10} more record(s) are available.")
                return "\n".join(lines)

    if re.search(r"\b(sum|total|average|avg|highest|lowest|max|min)\b", lower):
        query_tokens = _tokens(query)
        candidates: list[tuple[int, str, int, dict[str, Any]]] = []
        for record in records:
            for key, value in _flatten_numeric_fields(record).items():
                score = len(query_tokens & _tokens(key.replace("_", " ")))
                if score or key.lower() in lower:
                    candidates.append((score, key, value, record))
        if candidates:
            candidates.sort(key=lambda item: item[0], reverse=True)
            best_key = candidates[0][1]
            values = [(value, record) for _score, key, value, record in candidates if key == best_key]
            label = best_key.replace("_", " ")
            if re.search(r"\b(average|avg)\b", lower):
                return f"Average {label} is {_money_label(round(sum(value for value, _record in values) / len(values)))} across {len(values)} record(s)."
            if re.search(r"\b(highest|max)\b", lower):
                value, record = max(values, key=lambda item: item[0])
                return f"Highest {label} is {_money_label(value)} in {_record_summary(record)}."
            if re.search(r"\b(lowest|min)\b", lower):
                value, record = min(values, key=lambda item: item[0])
                return f"Lowest {label} is {_money_label(value)} in {_record_summary(record)}."
            return f"Total {label} is {_money_label(sum(value for value, _record in values))} across {len(values)} record(s)."

    if _is_general_detail_query(lower) or re.search(r"\b(summary|summarize|structure|schema|what type|document type)\b", lower):
        return _structured_overview(documents, records, include_examples=_is_general_detail_query(lower))

    return None


def _is_general_detail_query(lower: str) -> bool:
    return bool(re.fullmatch(r"\s*(give|show|tell me)?\s*(me\s+)?(more\s+)?details?\s*", lower))


def _structured_overview(documents: list[dict[str, Any]], records: list[dict[str, Any]], include_examples: bool = False) -> str:
    counts = Counter(record.get("kind", "unknown") for record in records)
    doc_types = Counter(doc.get("document_type", "unknown") for doc in documents)
    kind_parts = ", ".join(f"{kind}: {count}" for kind, count in sorted(counts.items())[:14])
    type_parts = ", ".join(f"{kind}: {count}" for kind, count in sorted(doc_types.items()))
    lines = [
        f"I have structured {len(documents)} document(s).",
        f"Document types: {type_parts or 'none'}.",
        f"Extracted record types: {kind_parts or 'none'}.",
    ]
    if include_examples:
        priority = {
            "agreement": 0,
            "renewal_action": 1,
            "contact": 2,
            "table_row": 3,
            "qa_pair": 4,
            "section": 5,
            "key_value": 6,
        }
        important = [
            record
            for record in records
            if record.get("kind") in {"agreement", "renewal_action", "contact", "table_row", "qa_pair", "key_value", "section"}
        ]
        important.sort(key=lambda record: priority.get(str(record.get("kind")), 99))
        important = important[:8]
        if important:
            lines.append("Key details:")
            for record in important:
                lines.append(f"- {_record_summary(record)}")
    return "\n".join(lines)


def _record_summary(record: dict[str, Any]) -> str:
    if record.get("kind") == "agreement":
        return ", ".join(_agreement_detail_parts(record))
    if record.get("kind") == "renewal_action":
        return (
            f"{record.get('agreement_id')}, {record.get('tenant')}, {record.get('property_id')}, "
            f"expires {record.get('expiry')}, proposed rent {_money_label(int(record.get('proposed_new_rent') or 0))}, "
            f"owner {record.get('action_owner')}"
        )
    if record.get("kind") == "contact":
        parts = [str(record.get("name")), str(record.get("contact_type"))]
        for field in ("phone", "email", "property_id", "budget", "requirement", "service", "assigned_to"):
            if record.get(field):
                parts.append(f"{field.replace('_', ' ')}: {record[field]}")
        return ", ".join(part for part in parts if part and part != "None")
    if record.get("kind") == "table_row" and isinstance(record.get("fields"), dict):
        return ", ".join(f"{key}: {_display_value(value)}" for key, value in list(record["fields"].items())[:6])
    for key in ("title", "text", "question", "key", "tenant", "name"):
        if record.get(key):
            if key == "key" and record.get("value"):
                return f"{record[key]}: {_display_value(record['value'])}"
            return str(record[key])[:240]
    visible = {
        key: value
        for key, value in record.items()
        if key not in {"source_file", "source_line"} and value not in (None, "")
    }
    return ", ".join(f"{key}: {_display_value(value)}" for key, value in list(visible.items())[:6])[:300]


def _display_value(value: Any) -> str:
    if isinstance(value, dict):
        if "raw" in value:
            return str(value["raw"])
        return json.dumps(value, ensure_ascii=False)
    return str(value)


def answer_structured_query(query: str, tenant_id: str | None) -> str | None:
    index = _read_index(tenant_id)
    documents = index.get("documents", [])
    records = [record for doc in documents for record in doc.get("records", [])]
    lower = query.lower()

    question_answer = _answer_question_from_records(query, records)
    if question_answer:
        return question_answer

    contact_answer = _answer_contact_query(query, records)
    if contact_answer:
        return contact_answer

    renewal_answer = _answer_renewal_query(query, records)
    if renewal_answer:
        return renewal_answer

    agreement_answer = _answer_agreement_query(query, records)
    if agreement_answer:
        return agreement_answer

    if re.search(r"\b(how many|count|number of|total)\b", lower) and re.search(r"\bdocuments?|files?\b", lower):
        counts = Counter(doc.get("document_type", "unknown") for doc in documents)
        parts = ", ".join(f"{kind}: {count}" for kind, count in sorted(counts.items()))
        return f"I have {len(documents)} structured document(s) in this workspace." + (f" Types: {parts}." if parts else "")

    properties = _dedupe_records([record for record in records if record.get("kind") == "property"], ("property_id",))
    if properties and re.search(r"\b(how many|count|number of|total)\b", lower) and re.search(r"\bproperties?|listings?\b", lower):
        status = next((item for item in STATUSES if item.lower() in lower), "")
        selected = [record for record in properties if not status or record.get("status") == status]
        residential = sum(1 for record in selected if record.get("category") == "residential")
        commercial = sum(1 for record in selected if record.get("category") == "commercial")
        ids = ", ".join(record.get("property_id", "") for record in selected if record.get("property_id"))
        label = f" {status.lower()}" if status else ""
        return f"You have {len(selected)}{label} propertie(s): {residential} residential and {commercial} commercial. Property IDs: {ids}."

    rent_records = _dedupe_records([record for record in records if record.get("kind") == "rent"], ("agreement_id",))
    if rent_records and re.search(r"\b(rent|tenant|agreement|paid|pending|overdue|balance|due)\b", lower):
        mentioned = [
            record
            for record in rent_records
            if any(
                str(record.get(field, "")).lower() in lower
                for field in ("agreement_id", "property_id")
                if record.get(field)
            )
            or any(part in lower for part in re.findall(r"[a-z]{3,}", str(record.get("tenant", "")).lower()))
        ]
        selected = mentioned or rent_records
        if re.search(r"\b(pending|overdue|balance|due|unpaid)\b", lower):
            pending = [record for record in selected if int(record.get("pending_amount") or 0) > 0]
            total = sum(int(record.get("pending_amount") or 0) for record in pending)
            if not pending:
                return "No pending rent was found in the structured rent records."
            lines = [f"Pending rent is Rs. {total:,} from {len(pending)} tenant/account(s):"]
            for record in pending:
                lines.append(
                    f"- {record.get('tenant')} ({record.get('property_id')}, {record.get('agreement_id')}): "
                    f"Rs. {int(record.get('pending_amount') or 0):,} pending "
                    f"(rent Rs. {int(record.get('rent_amount') or 0):,}, paid Rs. {int(record.get('paid_amount') or 0):,}, status: {record.get('status')})."
                )
            return "\n".join(lines)

    questions = _dedupe_records([record for record in records if record.get("kind") == "question"], ("source_file", "number", "title"))
    if questions and re.search(r"\b(how many|count|number of|total)\b", lower) and re.search(r"\bquestions?|problems?\b", lower):
        return f"I found {len(questions)} question(s) in the structured question records."

    qa_pair_answer = _answer_qa_pair_query(query, documents, records)
    if qa_pair_answer:
        return qa_pair_answer

    general_answer = _answer_general_structured_query(query, documents, records)
    if general_answer:
        return general_answer

    return None
