from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path

from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader

from utils.config import load_settings, tenant_generated_dir, tenant_uploads_dir


def _safe_stem(text: str) -> str:
    return re.sub(r"[^a-zA-Z0-9]+", "_", text).strip("_").lower()[:48] or "assistant_output"


def _generated_path(stem: str, suffix: str, tenant_id: str | None = None) -> Path:
    settings = load_settings()
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_dir = tenant_generated_dir(settings, tenant_id) if tenant_id else settings.generated_dir
    return output_dir / f"{stamp}_{_safe_stem(stem)}.{suffix}"


def create_docx(text: str, filename: str | None = None, tenant_id: str | None = None) -> Path:
    settings = load_settings()
    output_dir = tenant_generated_dir(settings, tenant_id) if tenant_id else settings.generated_dir
    path = output_dir / filename if filename else _generated_path(text, "docx", tenant_id)
    document = Document()
    for paragraph in text.splitlines() or [text]:
        document.add_paragraph(paragraph)
    document.save(path)
    return path


def create_text_file(text: str, stem: str = "assistant_output", tenant_id: str | None = None) -> Path:
    path = _generated_path(stem, "txt", tenant_id)
    path.write_text(text, encoding="utf-8")
    return path


def create_pdf(text: str, stem: str = "assistant_output", tenant_id: str | None = None) -> Path:
    path = _generated_path(stem, "pdf", tenant_id)
    lines = []
    for raw_line in text.splitlines():
        line = raw_line
        while len(line) > 88:
            lines.append(line[:88])
            line = line[88:]
        lines.append(line)

    pages = [lines[index : index + 48] for index in range(0, len(lines), 48)] or [[]]
    objects = [
        b"1 0 obj << /Type /Catalog /Pages 2 0 R >> endobj\n",
    ]
    page_ids = []
    font_id = 3 + (len(pages) * 2)
    for page_index, page_lines in enumerate(pages):
        page_id = 3 + (page_index * 2)
        content_id = page_id + 1
        page_ids.append(page_id)
        content = ["BT", "/F1 11 Tf", "50 780 Td", "14 TL"]
        for line_index, line in enumerate(page_lines):
            safe = line.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
            if line_index:
                content.append("T*")
            content.append(f"({safe}) Tj")
        content.append("ET")
        stream = "\n".join(content).encode("latin-1", errors="replace")
        objects.append(
            f"{page_id} 0 obj << /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Resources << /Font << /F1 {font_id} 0 R >> >> /Contents {content_id} 0 R >> endobj\n".encode("ascii")
        )
        objects.append(
            f"{content_id} 0 obj << /Length {len(stream)} >> stream\n".encode("ascii") + stream + b"\nendstream endobj\n"
        )
    kids = " ".join(f"{page_id} 0 R" for page_id in page_ids)
    objects.insert(1, f"2 0 obj << /Type /Pages /Kids [{kids}] /Count {len(page_ids)} >> endobj\n".encode("ascii"))
    objects.append(f"{font_id} 0 obj << /Type /Font /Subtype /Type1 /BaseFont /Helvetica >> endobj\n".encode("ascii"))
    out = [b"%PDF-1.4\n"]
    offsets = [0]
    for obj in objects:
        offsets.append(sum(len(part) for part in out))
        out.append(obj)
    xref_start = sum(len(part) for part in out)
    xref = [f"xref\n0 {len(objects) + 1}\n".encode("ascii"), b"0000000000 65535 f \n"]
    for offset in offsets[1:]:
        xref.append(f"{offset:010d} 00000 n \n".encode("ascii"))
    out.extend(xref)
    out.append(f"trailer << /Size {len(objects) + 1} /Root 1 0 R >>\n".encode("ascii"))
    out.append(f"startxref\n{xref_start}\n%%EOF".encode("ascii"))
    path.write_bytes(b"".join(out))
    return path


def create_document_bundle(text: str, stem: str = "assistant_output", tenant_id: str | None = None) -> list[Path]:
    return [create_text_file(text, stem, tenant_id), create_docx(text, tenant_id=tenant_id), create_pdf(text, stem, tenant_id)]


def _iter_docx_paragraphs(document: Document):
    for paragraph in document.paragraphs:
        yield paragraph
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def _docx_text(document: Document) -> str:
    return "\n".join(paragraph.text for paragraph in _iter_docx_paragraphs(document) if paragraph.text.strip())


def _replace_in_docx(document: Document, replacements: dict[str, str]) -> int:
    changed = 0
    for paragraph in _iter_docx_paragraphs(document):
        text = paragraph.text
        updated = text
        for old, new in replacements.items():
            updated = updated.replace(old, new)
        if updated != text:
            paragraph.text = updated
            changed += 1
    return changed


def _extract_party_names(instruction: str) -> dict[str, str]:
    parties: dict[str, str] = {}
    buyer_match = re.search(r"\bbuyer(?:\s+name)?\s*(?:is|=|:|as|to)?\s*(.+)", instruction, flags=re.I)
    seller_match = re.search(r"\b(?:seller|vendor)(?:\s+name)?\s*(?:is|=|:|as|to)?\s*(.+)", instruction, flags=re.I)

    if buyer_match:
        buyer = re.split(r"\s+(?:and\s+)?(?:seller|vendor)(?:\s+name)?\b|[,;]", buyer_match.group(1), maxsplit=1, flags=re.I)[0]
        buyer = re.sub(r"\b(for|in|on|with)\b.*$", "", buyer, flags=re.I).strip(" .,'-")
        if buyer:
            parties["buyer"] = buyer
    if seller_match:
        seller = re.split(r"\s+(?:and\s+)?(?:buyer|purchaser)(?:\s+name)?\b|[,;]", seller_match.group(1), maxsplit=1, flags=re.I)[0]
        seller = re.sub(r"\b(for|in|on|with)\b.*$", "", seller, flags=re.I).strip(" .,'-")
        if seller:
            parties["seller"] = seller
    return parties


def _party_replacements(parties: dict[str, str]) -> dict[str, str]:
    replacements: dict[str, str] = {}
    if buyer := parties.get("buyer"):
        replacements.update(
            {
                "[BUYER]": buyer,
                "[BUYER NAME]": buyer,
                "{{BUYER}}": buyer,
                "{{BUYER_NAME}}": buyer,
                "<BUYER>": buyer,
                "<BUYER NAME>": buyer,
                "BUYER_NAME": buyer,
                "Purchaser Name": buyer,
                "Buyer Name": buyer,
                "Name of Buyer": buyer,
            }
        )
    if seller := parties.get("seller"):
        replacements.update(
            {
                "[SELLER]": seller,
                "[SELLER NAME]": seller,
                "{{SELLER}}": seller,
                "{{SELLER_NAME}}": seller,
                "{{VENDOR}}": seller,
                "<SELLER>": seller,
                "<SELLER NAME>": seller,
                "SELLER_NAME": seller,
                "Vendor Name": seller,
                "Seller Name": seller,
                "Name of Seller": seller,
            }
        )
    return replacements


def _rewrite_docx_with_llm(path: Path, instruction: str, original_text: str, tenant_id: str | None = None) -> Path:
    from llm.ollama_client import generate_response

    prompt = (
        "Rewrite this DOCX template according to the user's instruction. "
        "Return only the revised document text. Preserve the legal/document structure, headings, clauses, "
        "numbering intent, and all unchanged details. Do not append an explanation.\n\n"
        f"Instruction:\n{instruction}"
    )
    revised = generate_response(prompt, f"Original DOCX text from {path.name}:\n{original_text[:90000]}")
    return create_docx(revised, f"{path.stem}_edited.docx", tenant_id)


def modify_docx(path: Path, instruction: str, tenant_id: str | None = None) -> Path:
    settings = load_settings()
    output_dir = tenant_generated_dir(settings, tenant_id) if tenant_id else settings.generated_dir
    document = Document(path)
    changed = 0
    replace_match = re.search(r"replace\s+(.+?)\s+with\s+(.+)", instruction, flags=re.I)
    if replace_match:
        old, new = replace_match.groups()
        changed += _replace_in_docx(document, {old: new})
    elif re.search(r"\b(add|append)\b", instruction, flags=re.I):
        content = re.sub(r"^\s*(add|append)\s+(paragraph|section|text)?\s*:?", "", instruction, flags=re.I).strip()
        document.add_paragraph(content or instruction)
        changed += 1
    else:
        parties = _extract_party_names(instruction)
        if parties:
            changed += _replace_in_docx(document, _party_replacements(parties))
        if changed == 0:
            return _rewrite_docx_with_llm(path, instruction, _docx_text(document), tenant_id)
    output = output_dir / f"{path.stem}_edited{path.suffix}"
    document.save(output)
    return output


def update_xlsx(path: Path, instruction: str, tenant_id: str | None = None) -> Path:
    settings = load_settings()
    output_dir = tenant_generated_dir(settings, tenant_id) if tenant_id else settings.generated_dir
    workbook = load_workbook(path)
    sheet = workbook.active
    lower = instruction.lower()

    percent_match = re.search(r"column\s+([a-z0-9_ ]+)\s*\+(\d+(?:\.\d+)?)%", lower)
    if percent_match:
        header_name = percent_match.group(1).strip()
        multiplier = 1 + float(percent_match.group(2)) / 100
        headers = {
            str(cell.value).strip().lower(): cell.column
            for cell in sheet[1]
            if cell.value is not None
        }
        column = headers.get(header_name)
        if column:
            for row in range(2, sheet.max_row + 1):
                cell = sheet.cell(row=row, column=column)
                if isinstance(cell.value, (int, float)):
                    cell.value = cell.value * multiplier

    formula_match = re.search(r"formula\s+([a-z]+\d+)\s*=\s*(.+)", instruction, flags=re.I)
    if formula_match:
        cell_ref, formula = formula_match.groups()
        sheet[cell_ref.upper()] = formula if formula.startswith("=") else f"={formula}"

    set_match = re.search(r"(?:set|update)\s+([a-z]+\d+)\s*(?:=|to)\s*(.+)", instruction, flags=re.I)
    if set_match:
        cell_ref, value = set_match.groups()
        value = value.strip()
        try:
            parsed_value = float(value) if "." in value else int(value)
        except ValueError:
            parsed_value = value
        sheet[cell_ref.upper()] = parsed_value

    status_match = re.search(r"mark\s+(?:all\s+)?(?:rows\s+)?(?:as\s+)?([a-zA-Z ]+)", instruction, flags=re.I)
    if status_match:
        headers = [str(cell.value or "").strip().lower() for cell in sheet[1]]
        if "status" in headers:
            status_col = headers.index("status") + 1
        else:
            status_col = sheet.max_column + 1
            sheet.cell(row=1, column=status_col).value = "status"
        status_value = status_match.group(1).strip().title()
        for row in range(2, sheet.max_row + 1):
            sheet.cell(row=row, column=status_col).value = status_value

    output = output_dir / f"{path.stem}_edited{path.suffix}"
    workbook.save(output)
    return output


def edit_pdf_as_new_pdf(path: Path, instruction: str, revised_text: str | None = None, tenant_id: str | None = None) -> Path:
    original = read_pdf_text(path, max_chars=90000)
    if revised_text:
        text = revised_text
    else:
        text = f"Edited PDF content based on instruction: {instruction}\n\nOriginal extracted text:\n{original}"
    return create_pdf(text, f"{path.stem}_edited", tenant_id)


def _money_to_int(value: str) -> int:
    cleaned = re.sub(r"[^\d]", "", value)
    return int(cleaned or "0")


def _format_money(value: int) -> str:
    return f"{value:,}"


def _find_rent_tracker_pdf(tenant_id: str | None = None) -> Path | None:
    settings = load_settings()
    root = tenant_uploads_dir(settings, tenant_id) if tenant_id else settings.uploads_dir
    candidates = []
    if not root.exists():
        return None
    for path in root.rglob("*.pdf"):
        name = path.name.lower()
        if "rent" in name and ("collection" in name or "tracker" in name):
            candidates.append(path)
    return sorted(candidates, key=lambda item: item.stat().st_mtime, reverse=True)[0] if candidates else None


def _payment_date(instruction: str) -> str:
    date_match = re.search(r"\b(\d{1,2}[-/][A-Za-z]{3}[-/]\d{2,4}|\d{1,2}[-/]\d{1,2}[-/]\d{2,4})\b", instruction)
    if date_match:
        return date_match.group(1).replace("/", "-")
    return datetime.now().strftime("%d-%b-%y")


def _payment_mode(instruction: str) -> str:
    for mode in ("NEFT", "UPI", "Cash", "Cheque", "RTGS", "IMPS"):
        if re.search(rf"\b{mode}\b", instruction, flags=re.I):
            return mode
    return "Recorded"


def _paid_amount_from_instruction(instruction: str, rent_amount: int) -> int:
    amount_match = re.search(r"(?:paid|received|payment(?:\s+of)?)\s+(?:rs\.?|inr)?\s*(\d[\d,]*)", instruction, flags=re.I)
    if amount_match:
        return _money_to_int(amount_match.group(1))
    if re.search(r"\b(full|complete|cleared|paid\s+(?:its|his|her|their)?\s*rent)\b", instruction, flags=re.I):
        return rent_amount
    return rent_amount


def update_rent_tracker_pdf(instruction: str, tenant_id: str | None = None) -> tuple[Path, str] | None:
    agreement_match = re.search(r"\b([RC]A-\d{3,})\b", instruction, flags=re.I)
    has_payment_word = re.search(r"\b(paid|received|cleared|settled)\b", instruction, flags=re.I)
    has_action_word = re.search(r"\b(update|mark|record|change|set|edit|modify)\b", instruction, flags=re.I)
    looks_like_question = instruction.strip().endswith("?") or re.match(r"\s*(has|does|did|is|was|what|who|how)\b", instruction, flags=re.I)
    if not agreement_match or not has_payment_word or (looks_like_question and not has_action_word):
        return None
    if not has_action_word and not re.search(r"\b(?:this\s+tenant|tenant|rent)\s+(?:has\s+)?(?:paid|cleared|settled)\b|\bpaid\s+(?:its|his|her|their)?\s*rent\b", instruction, flags=re.I):
        return None

    agreement_id = agreement_match.group(1).upper()
    source = _find_rent_tracker_pdf(tenant_id)
    if not source:
        raise FileNotFoundError("I could not find a rent collection tracker PDF in workspace/uploads.")

    text = read_pdf_text(source, max_chars=50000)
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    row_index = next((index for index, line in enumerate(lines) if line.upper() == agreement_id), None)
    if row_index is None:
        raise ValueError(f"I could not find {agreement_id} in {source.name}.")

    is_residential = agreement_id.startswith("RA-")
    required_length = 9 if is_residential else 8
    if row_index + required_length > len(lines):
        raise ValueError(f"I found {agreement_id}, but its table row is incomplete.")

    tenant = lines[row_index + 1]
    property_id = lines[row_index + 2]
    rent_amount = _money_to_int(lines[row_index + 4])
    paid_amount = min(_paid_amount_from_instruction(instruction, rent_amount), rent_amount)
    paid_on = _payment_date(instruction)
    mode = _payment_mode(instruction)
    status = "Paid" if paid_amount >= rent_amount else "Partial" if paid_amount > 0 else "Overdue"

    lines[row_index + 5] = _format_money(paid_amount)
    lines[row_index + 6] = paid_on
    if is_residential:
        lines[row_index + 7] = mode
        lines[row_index + 8] = status
    else:
        lines[row_index + 7] = status

    lines.append("")
    lines.append("Update Log")
    lines.append(
        f"{datetime.now().strftime('%d-%b-%Y %H:%M')}: {agreement_id} ({tenant}, {property_id}) marked {status}. "
        f"Rent Rs. {_format_money(rent_amount)}, paid Rs. {_format_money(paid_amount)}, balance Rs. {_format_money(max(rent_amount - paid_amount, 0))}."
    )

    output = create_pdf("\n".join(lines), f"{source.stem}_updated", tenant_id)
    summary = (
        f"Updated {source.name}: {agreement_id} ({tenant}, {property_id}) marked {status}. "
        f"Paid Rs. {_format_money(paid_amount)} against rent Rs. {_format_money(rent_amount)}."
    )
    return output, summary


def process_uploaded_document(path: Path, instruction: str, revised_text: str | None = None, tenant_id: str | None = None) -> Path | None:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return modify_docx(path, instruction, tenant_id)
    if suffix == ".xlsx":
        return update_xlsx(path, instruction, tenant_id)
    if suffix == ".pdf":
        return edit_pdf_as_new_pdf(path, instruction, revised_text, tenant_id)
    return None


def read_pdf_text(path: Path, max_chars: int = 90000) -> str:
    reader = PdfReader(path)
    parts = []
    for page in reader.pages:
        parts.append(page.extract_text() or "")
        if sum(len(part) for part in parts) >= max_chars:
            break
    return "\n".join(parts).strip()[:max_chars]
